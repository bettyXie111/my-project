from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
MATERIALS = ROOT / "copyright-application-materials"


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    document = Document()
    lines = md_path.read_text(encoding="utf-8").splitlines()
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("#### "):
            document.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped and stripped[0].isdigit() and ". " in stripped:
            _, content = stripped.split(". ", 1)
            document.add_paragraph(content.strip(), style="List Number")
        else:
            document.add_paragraph(line)
    document.save(docx_path)


def main() -> None:
    pairs = [
        ("软件著作权登记申请表.md", "软件著作权登记申请表.docx"),
        ("源代码文档.md", "源代码文档.docx"),
        ("操作手册.md", "操作手册.docx"),
    ]
    for md_name, docx_name in pairs:
        md_path = MATERIALS / md_name
        docx_path = MATERIALS / docx_name
        md_to_docx(md_path, docx_path)
        print(f"generated: {docx_path}")


if __name__ == "__main__":
    main()
