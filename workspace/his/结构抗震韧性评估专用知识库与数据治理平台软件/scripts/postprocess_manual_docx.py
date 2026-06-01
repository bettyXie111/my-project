from pathlib import Path
from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def main():
    root = Path(__file__).resolve().parents[1]
    docx_path = root / "copyright-application-materials" / "企业运营全流程数字化管控平台软件操作手册.docx"
    doc = Document(str(docx_path))

    # insert cover page
    title = "企业运营全流程数字化管控平台软件"
    p = doc.paragraphs[0].insert_paragraph_before()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = doc.styles["Title"].font.size
    p2 = p.insert_paragraph_before()
    p2.add_run("")
    p3 = p2.insert_paragraph_before()
    p3.add_run("")
    p.runs[0].add_break()  # line break in cover

    # page break after cover
    pb = p.insert_paragraph_before()
    pb_run = pb.add_run()
    pb_run.add_break(WD_BREAK.PAGE)

    # header + page number
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = f"{title}V1.0"
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    add_page_number(footer)

    doc.save(str(docx_path))
    print(docx_path)


if __name__ == "__main__":
    main()
