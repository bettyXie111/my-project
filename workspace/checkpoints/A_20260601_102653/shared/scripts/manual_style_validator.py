# -*- coding: utf-8 -*-
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SKILL_PATH = Path(r"C:\Users\31556\.codex\skills\custom\chinese-copyright-application\SKILL.md")
MANUAL_TEMPLATE_PATH = Path(r"C:\Users\31556\.codex\skills\custom\chinese-copyright-application\references\operation-manual-template.docx")
WORKSPACE_RULE_FILES = (
    "CodeDocumentGenerationRules.md",
    "AutomaticVerificationRulesFailureDetermination.md",
)


def _find_workspace_root(start: Path) -> Path:
    for parent in [start, *start.parents]:
        if (parent / "pipeline_guard.py").exists() and (parent / "shared" / "scripts").exists():
            return parent
    raise FileNotFoundError(f"Cannot locate workspace root from: {start}")


def _read_workspace_rule_files() -> dict[str, str]:
    root = _find_workspace_root(Path(__file__).resolve())
    out: dict[str, str] = {}
    for name in WORKSPACE_RULE_FILES:
        if name == "CodeDocumentGenerationRules.md":
            path = root / "shared" / "docs" / "rules" / name
        else:
            path = root / "shared" / "docs" / "rules" / name
        if not path.exists():
            raise FileNotFoundError(f"Workspace rule file missing: {path}")
        out[name] = path.read_text(encoding="utf-8", errors="strict")
    print(f"[manual_style_validator] loaded workspace rules: {', '.join(sorted(out.keys()))}")
    return out


def get_manual_template_path() -> Path:
    if not MANUAL_TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Cannot locate manual template: {MANUAL_TEMPLATE_PATH}")
    return MANUAL_TEMPLATE_PATH


def assert_skill_has_cover_paging_rule() -> None:
    if not SKILL_PATH.exists():
        raise FileNotFoundError(f"Cannot locate skill file: {SKILL_PATH}")
    text = SKILL_PATH.read_text(encoding="utf-8")
    if "封面后分页进入正文" not in text:
        raise ValueError("Skill 缺少“封面后分页进入正文”规则。")


def _is_step_line(text: str) -> bool:
    if not re.match(r"^\d+\.\s+", text):
        return False
    body = re.sub(r"^\d+\.\s+", "", text).strip()
    if body.endswith(("。", "；", "！", "!")):
        return True
    if len(body) >= 12:
        return True
    return bool(
        re.match(
            r"^(打开|进入|点击|输入|选择|查看|提交|确认|返回|若|如果|然后|再|并|在|先|后|将|对|处置|过滤|识别|核查)",
            body,
        )
    )


def normalize_manual_markdown_structure(src_lines: list[str]) -> list[str]:
    cn_map = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "七", 8: "八", 9: "九", 10: "十", 11: "十一", 12: "十二"}
    out: list[str] = []
    for line in src_lines:
        t = line.strip()
        m1 = re.match(r"^(\d+)\.\s+(.+)$", t)
        if m1 and not _is_step_line(t):
            idx = int(m1.group(1))
            title = m1.group(2).strip()
            out.append(f"## {cn_map.get(idx, str(idx))}、{title}")
            out.append("")
            continue
        m2 = re.match(r"^(\d+\.\d+)\s+(.+)$", t)
        if m2:
            out.append(f"### {m2.group(1)} {m2.group(2).strip()}")
            out.append("")
            continue
        out.append(line)
    return out


def enforce_manual_cover(docx_path: Path, requirement_name: str) -> None:
    doc = Document(str(docx_path))
    expected_title = f"{requirement_name} V1.0"
    expected_sub = "操作手册"
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    if len(non_empty) < 2:
        first = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        title_p = first.insert_paragraph_before(expected_title)
        sub_p = first.insert_paragraph_before(expected_sub)
    else:
        title_p, sub_p = non_empty[0], non_empty[1]
        title_p.text = expected_title
        sub_p.text = expected_sub

    idx_title = None
    idx_sub = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if idx_title is None and t == expected_title:
            idx_title = i
            continue
        if idx_title is not None and t == expected_sub:
            idx_sub = i
            break
    if idx_title is None or idx_sub is None:
        idx_title, idx_sub = 0, 1
    if idx_sub - idx_title > 1:
        for i in range(idx_sub - 1, idx_title, -1):
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    for p in (title_p, sub_p):
        p.style = doc.styles["Normal"]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if not p.runs:
            p.add_run(p.text)
        for run in p.runs:
            run.font.name = "宋体"
            run.font.size = Pt(24)

    section = doc.sections[0]
    sect_pr = section._sectPr
    for v in sect_pr.findall(qn("w:vAlign")):
        sect_pr.remove(v)
    v_align = OxmlElement("w:vAlign")
    # Apply vertical alignment for the whole section. Keep it as top so body pages
    # are not vertically centered (which can create large top blanks on pages with
    # little content). Cover centering is achieved via a bounded space_before.
    v_align.set(qn("w:val"), "top")
    sect_pr.append(v_align)
    # Cover vertical centering: keep section vAlign=top (so body pages are normal),
    # and use a conservative-but-center-ish spacing on the title line.
    # A4 usable height is ~700pt after margins; 200pt places the 2-line 24pt block
    # close to mid-page without pushing it out of view.
    title_p.paragraph_format.space_before = Pt(200)

    if not sub_p.runs:
        sub_p.add_run("")
    sub_p.runs[-1].add_break(WD_BREAK.PAGE)
    passed_sub = False
    for p in list(doc.paragraphs):
        if not passed_sub:
            if p is sub_p:
                passed_sub = True
            continue
        if p.text.strip():
            p.paragraph_format.space_before = Pt(0)
            break
        p._element.getparent().remove(p._element)

    doc.save(str(docx_path))


def cleanup_literal_newpage_markers(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    for p in list(doc.paragraphs):
        if p.text.strip().replace("\\", "").replace(" ", "").lower() == "newpage":
            p._element.getparent().remove(p._element)
    doc.save(str(docx_path))


def normalize_body_pagination_spacing(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    passed_cover = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if not passed_cover:
            if t == "操作手册":
                passed_cover = True
            continue
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.page_break_before = False
        pf.keep_with_next = False
        pf.keep_together = False
    used_style_ids = {p.style.style_id for p in doc.paragraphs if p.text.strip()}
    for st in doc.styles:
        if st.style_id not in used_style_ids:
            continue
        try:
            st.paragraph_format.space_before = Pt(0)
            st.paragraph_format.page_break_before = False
            st.paragraph_format.keep_with_next = False
            st.paragraph_format.keep_together = False
        except Exception:
            continue
    doc.save(str(docx_path))


def assert_manual_cover_rules(docx_path: Path, requirement_name: str) -> None:
    doc = Document(str(docx_path))
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    if len(non_empty) < 2:
        raise ValueError("操作手册封面校验失败：缺少封面双行。")
    if non_empty[0].text.strip() != f"{requirement_name} V1.0" or non_empty[1].text.strip() != "操作手册":
        raise ValueError("操作手册封面校验失败：封面双行不匹配。")
    if non_empty[0].style.name != "Normal" or non_empty[1].style.name != "Normal":
        raise ValueError("操作手册封面校验失败：封面双行样式必须为 Normal，禁止使用 Heading 样式。")
    if non_empty[0].alignment != WD_PARAGRAPH_ALIGNMENT.CENTER or non_empty[1].alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
        raise ValueError("操作手册封面校验失败：封面双行必须水平居中。")
    title_idx = -1
    sub_idx = -1
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if title_idx < 0 and txt == f"{requirement_name} V1.0":
            title_idx = i
            continue
        if title_idx >= 0 and txt == "操作手册":
            sub_idx = i
            break
    if title_idx >= 0 and sub_idx > title_idx:
        if any(p.text.strip() == "" for p in doc.paragraphs[title_idx + 1 : sub_idx]):
            raise ValueError("操作手册封面校验失败：主标题与副标题之间存在空行。")
    has_page_break = False
    for br in non_empty[1]._p.findall(".//w:br", {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
        if br.get(qn("w:type")) == "page":
            has_page_break = True
            break
    if not has_page_break:
        raise ValueError("操作手册封面校验失败：副标题后缺少分页符。")

    body_first = None
    passed_cover = False
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not passed_cover:
            if txt == "操作手册":
                passed_cover = True
            continue
        if txt:
            body_first = txt
            break
    if not body_first:
        raise ValueError("操作手册正文校验失败：封面后正文首行为空。")


def assert_manual_markdown_heading_safety(manual_md_path: Path) -> None:
    lines = manual_md_path.read_text(encoding="utf-8", errors="ignore").splitlines()
    bad: list[str] = []
    for i, line in enumerate(lines, start=1):
        t = line.strip()
        if re.match(r"^##\s+[一二三四五六七八九十]+、", t) and re.search(r"打开|进入|点击|输入|选择|查看|提交|确认|返回|若|如果", t):
            bad.append(f"L{i}:{t}")
    if bad:
        raise ValueError("操作手册标题校验失败：检测到步骤语句被错误提升为一级标题。")


def assert_manual_ui_screenshot_interaction_rules(manual_md_path: Path) -> None:
    _ = _read_workspace_rule_files()
    text = manual_md_path.read_text(encoding="utf-8", errors="ignore")

    # Hard gate: forbid modal/dialog keywords in headings to avoid Word目录/标题污染。
    forbidden = ("弹窗", "对话框", "确认弹窗")
    bad_headings: list[str] = []
    for i, line in enumerate(text.splitlines(), start=1):
        t = line.strip()
        if not re.match(r"^#{2,4}\s+", t):
            continue
        if any(k in t for k in forbidden):
            bad_headings.append(f"L{i}:{t}")
    if bad_headings:
        preview = "\n".join(bad_headings[:20])
        raise ValueError(
            "操作手册标题校验失败：标题行禁止包含“弹窗/对话框/确认弹窗”等关键词，请修改标题后再生成。\n"
            + preview
        )

    required_ops = ("新增", "编辑", "删除", "查看")
    missing = [op for op in required_ops if op not in text]
    if missing:
        raise ValueError("操作手册截图/交互校验失败：未体现关键操作按钮文案：" + "、".join(missing))

    modal_hits = len(re.findall(r"弹窗|对话框|确认(?!码)|确认删除|删除确认", text))
    if modal_hits < 2:
        raise ValueError("操作手册截图/交互校验失败：未满足“至少 2 张弹窗/确认类截图说明”的要求（文档中相关描述不足）。")


def assert_manual_template_skeleton(manual_md_path: Path, requirement_name: str) -> None:
    """
    Ensure the operation manual markdown follows the official template skeleton.

    We deliberately validate only the *chapter-level* structure and key “图题” rules
    to avoid forcing projects to implement UI widgets they do not have yet.
    """
    template_md = Path(
        r"C:\Users\31556\.codex\skills\custom\chinese-copyright-application\references\operation-manual-template.md"
    )
    if not template_md.exists():
        raise FileNotFoundError(f"Cannot locate manual template md: {template_md}")

    text = manual_md_path.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()

    # 1) Manual markdown title line policy (skill-aligned):
    # Markdown 正文首行不得再写合并标题“软件全称 V1.0 操作手册”，避免与 build 阶段注入的封面/标题冲突。
    first_non_empty = next((l.strip() for l in lines if l.strip()), "")
    if first_non_empty.startswith("# "):
        merged = first_non_empty[2:].strip().replace(" ", "")
        if requirement_name.replace(" ", "") in merged and "V1.0" in merged and "操作手册" in merged:
            raise ValueError("操作手册模板校验失败：正文 md 首行不得写“软件全称 V1.0 操作手册”合并标题（封面由生成阶段注入）。")

    # 2) Chapter-level skeleton (as in template).
    # Accept chapter numbering shifts (some projects add “文档定位与适用对象” as the first chapter),
    # but the template chapter *titles* must exist.
    required_chapter_titles = [
        "系统入口与登录",
        "首页操作",
        # Accept variants when a project has no “二级菜单展开/下拉框/复选框” yet.
        ("一级功能模块与二级菜单操作", "一级功能模块"),
        "端到端操作流程演示",
        "附录",
    ]
    missing: list[str] = []
    for item in required_chapter_titles:
        if isinstance(item, tuple):
            if not any(token in text for token in item):
                missing.append(item[0])
            continue
        if item not in text:
            missing.append(item)
    if missing:
        raise ValueError("操作手册模板校验失败：缺少模板章节：" + "、".join(missing))

    # 3) Every screenshot must have a nearby caption line that starts with “图”.
    # This matches the template’s “图片 + 图题” pattern.
    for i, line in enumerate(lines):
        m = re.search(r"!\[[^\]]*\]\(images/([^)]+)\)", line)
        if not m:
            continue
        window = "\n".join(lines[i : min(i + 12, len(lines))])
        if not re.search(r"^图\d+-\d+[a-zA-Z]?\s+.+", window, flags=re.MULTILINE):
            raise ValueError(f"操作手册模板校验失败：图片缺少图题说明：{m.group(1)}")
