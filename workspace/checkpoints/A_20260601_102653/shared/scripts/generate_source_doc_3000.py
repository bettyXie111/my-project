from pathlib import Path
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def main():
    root = Path(__file__).resolve().parents[1]
    md = root / "copyright-application-materials" / "源代码文档.md"
    out = root / "copyright-application-materials" / "源代码文档-最新.docx"
    lines = md.read_text(encoding="utf-8").splitlines()

    # remove existing numbering and title, keep pure code flow
    clean = []
    import re
    for ln in lines:
        m = re.match(r"^\s*\d+\.\s+(.*)$", ln)
        txt = m.group(1) if m else ln
        txt = txt.lstrip("\ufeff")
        if txt.strip() in {"# 源代码文档", "源代码文档"}:
            continue
        clean.append(txt)

    while len(clean) < 3000:
        clean.append("")
    clean = clean[:3000]

    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "企业运营全流程数字化管控平台软件V1.0"
    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    add_page_number(footer)

    for i, content in enumerate(clean, start=1):
        text = content if content.strip() else " "
        p = doc.add_paragraph(text)
        pf = p.paragraph_format
        pf.space_before = 0
        pf.space_after = 0
        pf.line_spacing = 1.0
        if i % 50 == 0 and i != 3000:
            p.runs[-1].add_break(WD_BREAK.PAGE)

    doc.save(str(out))
    print(out)


if __name__ == "__main__":
    main()
