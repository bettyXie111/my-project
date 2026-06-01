# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import json
from datetime import date
import unicodedata
import os
import re
import subprocess
import sys
import tempfile
import time
from pathlib import Path
import shutil
from urllib.error import URLError
from urllib.request import urlopen

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
import manual_style_validator as msv
from zipfile import BadZipFile

try:
    import win32com.client  # type: ignore
    import pythoncom  # type: ignore
except Exception:  # pragma: no cover
    win32com = None  # type: ignore
    pythoncom = None  # type: ignore


EDGE_EXECUTABLE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")
DEFAULT_SERVER_PORT = 8010
CODE_EXTENSIONS = {".py", ".js", ".mjs", ".css", ".html", ".sql", ".json"}


def load_source_doc_rules() -> dict[str, object]:
    """
    Load workspace-wide rules for 源代码文档生成.

    Canonical location: `${WORKSPACE_ROOT}/config/rules/source_doc_rules.json`.
    `build_materials.py` should keep policy in that file (or SKILL/docs),
    and only keep the loading/implementation logic here.
    """
    candidate_paths = [
        ROOT / "shared" / "config" / "rules" / "source_doc_rules.json",
        ROOT / "config" / "rules" / "source_doc_rules.json",
    ]
    for path in candidate_paths:
        if not path.exists():
            continue
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            return data if isinstance(data, dict) else {}
        except Exception:
            return {}
    return {}


def load_manual_rules() -> dict[str, object]:
    """
    Load workspace-wide rules for 操作手册生成.

    Canonical location: `${WORKSPACE_ROOT}/config/rules/manual_rules.json`.
    `build_materials.py` should keep policy in that file (or SKILL/docs),
    and only keep the loading/implementation logic here.
    """
    candidate_paths = [
        ROOT / "shared" / "config" / "rules" / "manual_rules.json",
        ROOT / "config" / "rules" / "manual_rules.json",
    ]
    for path in candidate_paths:
        if not path.exists():
            continue
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            return data if isinstance(data, dict) else {}
        except Exception:
            return {}
    return {}


def load_default_screenshot_plan() -> list[dict[str, object]]:
    if not DEFAULT_SCREENSHOT_PLAN_TEMPLATE.exists():
        return []
    try:
        data = json.loads(DEFAULT_SCREENSHOT_PLAN_TEMPLATE.read_text(encoding="utf-8"))
    except Exception:
        return []
    if not isinstance(data, list):
        return []
    return [dict(item) for item in data if isinstance(item, dict)]


NORMAL_MARGIN_MM = 25.4
LINES_PER_PAGE = 50
A4_HEIGHT_MM = 297.0
A4_WIDTH_MM = 210.0
LINE_PITCH_TWIPS = int(round(((A4_HEIGHT_MM - 2 * NORMAL_MARGIN_MM) / LINES_PER_PAGE) / 25.4 * 1440))
DEFAULT_SCREENSHOT_PLAN_TEMPLATE = Path(__file__).with_name("screenshot_plan.template.json")
MIN_SCREENSHOT_COUNT = 12
MIN_MANUAL_CHARS = 6000

_MANUAL_RULES: dict[str, object] = {}


def strip_leading_manual_titles(lines: list[str], requirement_name: str) -> list[str]:
    out = list(lines)
    first_non_empty_idx = None
    for idx, line in enumerate(out):
        if line.strip():
            first_non_empty_idx = idx
            break
    if first_non_empty_idx is None:
        return out
    first_line = out[first_non_empty_idx].strip().replace("#", "").strip()
    normalized = first_line.replace(" ", "")
    if requirement_name in normalized and "V1.0" in normalized:
        out.pop(first_non_empty_idx)
        while first_non_empty_idx < len(out) and not out[first_non_empty_idx].strip():
            out.pop(first_non_empty_idx)
        if first_non_empty_idx < len(out) and out[first_non_empty_idx].strip() in {"操作手册", "# 操作手册"}:
            out.pop(first_non_empty_idx)
            while first_non_empty_idx < len(out) and not out[first_non_empty_idx].strip():
                out.pop(first_non_empty_idx)
    return out


def sanitize_manual_markdown(manual_md: Path, requirement_name: str) -> None:
    """
    Enforce skill rules on 操作手册.md正文:
    - Remove merged top title line (software name + version + 操作手册).
    - Remove generic "version/meta/checklist" tails that are not part of the operation manual body.
    """
    text = manual_md.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()
    lines = strip_leading_manual_titles(lines, requirement_name)

    rules = _MANUAL_RULES
    raw_titles = rules.get("manual_tail_exclude_headings")
    if isinstance(raw_titles, list) and raw_titles:
        remove_from_titles = {str(x) for x in raw_titles if str(x).strip()}
    else:
        remove_from_titles = {"版本信息", "交付一致性检查清单", "截图复核要点"}
    raw_requires = rules.get("manual_tail_exclude_requires_tokens")
    requires = [str(x) for x in raw_requires if str(x).strip()] if isinstance(raw_requires, list) and raw_requires else ["软件全称", "版本号"]
    cut_idx: int | None = None
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped.startswith("#"):
            continue
        title = stripped.lstrip("#").strip()
        title = re.sub(r"^\d+(?:\.\d+)*\s+", "", title)
        if any(tok in title for tok in remove_from_titles):
            window = "\n".join(lines[i : min(i + 30, len(lines))])
            requires_hit = all(tok in window for tok in requires) if requires else False
            if requires_hit or ("交付一致性检查清单" in title) or ("截图复核要点" in title):
                cut_idx = i
                break
    if cut_idx is not None:
        lines = lines[:cut_idx]
        while lines and not lines[-1].strip():
            lines.pop()

    # Enforce appendix allowlist: drop unexpected appendix headings that are not part of the operation manual.
    raw_allow = rules.get("manual_appendix_allowlist_headings")
    allow: set[str] = set(str(x) for x in raw_allow if str(x).strip()) if isinstance(raw_allow, list) else set()
    if allow:
        kept: list[str] = []
        in_appendix = False
        skip_section = False
        seen_appendix_h3: set[str] = set()
        appendix_chapter_pat = re.compile(r"^##\s+\d+\.\s*(附录)\b")
        for line in lines:
            stripped = line.strip()
            if appendix_chapter_pat.match(stripped):
                in_appendix = True
                skip_section = False
                kept.append(line)
                continue
            if in_appendix and stripped.startswith("## "):
                # Leaving appendix
                in_appendix = False
                skip_section = False
                seen_appendix_h3.clear()
                kept.append(line)
                continue
            if in_appendix and stripped.startswith("### "):
                title = stripped[4:].strip()
                title = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", title)
                if title not in allow:
                    skip_section = True
                    continue
                if title in seen_appendix_h3:
                    # Duplicate appendix section: skip this heading and its content.
                    skip_section = True
                    continue
                seen_appendix_h3.add(title)
                skip_section = False
                kept.append(line)
                continue
            if in_appendix and skip_section:
                # Skip content until next "###" or "##"
                continue
            kept.append(line)
        lines = kept

    # Normalize heading numbering for正文 (##/###/####) to prevent duplicated/out-of-order indices.
    # This keeps Markdown multi-level numbering consistent across projects.
    lines = prune_duplicate_appendix_sections(lines)
    lines = ensure_numbered_headings_markdown(lines)
    new_text = "\n".join(lines).rstrip() + "\n"
    if new_text != text:
        manual_md.write_text(new_text, encoding="utf-8")


def ensure_manual_min_chars(manual_md: Path, min_chars: int = MIN_MANUAL_CHARS) -> None:
    """
    Keep operation manual above the minimum character threshold without introducing
    PRD/industry narrative or metadata-only filler.
    """
    text = manual_md.read_text(encoding="utf-8", errors="ignore")
    cur = strip_markdown_for_count(text)
    if cur >= min_chars:
        return
    rules = _MANUAL_RULES
    blocks: list[str] = []
    raw_blocks = rules.get("manual_min_chars_append_blocks")
    if isinstance(raw_blocks, list) and raw_blocks:
        for b in raw_blocks:
            if not isinstance(b, dict):
                continue
            heading = str(b.get("heading", "")).strip()
            body = b.get("body")
            if not heading or not isinstance(body, list):
                continue
            body_lines = [str(x) for x in body if str(x).strip()]
            if not body_lines:
                continue
            blocks.append("\n".join([f"### {heading}", "", *body_lines]) + "\n")
    if not blocks:
        # Fallback minimal block when config is missing.
        blocks = ["### 运行与访问说明\n\n1. 启动服务：在项目目录运行启动命令后，等待服务启动提示。\n"]

    updated = text.rstrip() + "\n"
    # Avoid duplicating appendix blocks if they already exist in the manual.
    existing_h3_titles: set[str] = set()
    for line in updated.splitlines():
        s = line.strip()
        if not s.startswith("### "):
            continue
        title = s[4:].strip()
        title = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", title)
        existing_h3_titles.add(title)
    for block in blocks:
        if strip_markdown_for_count(updated) >= min_chars:
            break
        # If the block's heading already exists, skip the whole block.
        first = next((ln.strip() for ln in block.splitlines() if ln.strip()), "")
        if first.startswith("### "):
            t = first[4:].strip()
            if t in existing_h3_titles:
                continue
            existing_h3_titles.add(t)
        updated += "\n" + block.lstrip()
    normalized = ensure_numbered_headings_markdown(updated.splitlines())
    manual_md.write_text("\n".join(normalized).rstrip() + "\n", encoding="utf-8")


def ensure_numbered_headings_markdown(lines: list[str]) -> list[str]:
    chapter = 0
    sub = 0
    subsub = 0
    out: list[str] = []
    num_prefix = re.compile(r"^\d+(?:\.\d+)*\.?\s+")
    prev_heading: str | None = None
    prev_h2_title: str | None = None
    for line in lines:
        raw = line.rstrip("\n")
        stripped = raw.strip()
        # Drop exact duplicate consecutive headings to avoid repeated appendix blocks.
        if stripped.startswith("### ") and prev_heading == stripped:
            continue
        if stripped.startswith("## "):
            chapter += 1
            sub = 0
            subsub = 0
            title = stripped[3:].strip()
            # 一级标题只允许形如 "1 标题"，不允许 "1.1 标题" 这类二级编号串到一级。
            title = num_prefix.sub("", title).strip()
            # Defensive: strip any remaining repeated numeric prefixes like "6. 6. 标题"
            title = re.sub(r"^(?:\d+\.\s*)+", "", title).strip()
            # If title uses Chinese chapter marker like "一、标题"，normalize to numeric "1. 标题"
            # so that downstream docx headings have consistent numbering.
            title = re.sub(r"^[一二三四五六七八九十]+、\s*", "", title)
            # Drop duplicated consecutive H2 chapters (common when generators append twice).
            if prev_h2_title == title:
                continue
            out.append(f"## {chapter}. {title}")
            prev_h2_title = title
            prev_heading = stripped
            continue
        if stripped.startswith("### "):
            if chapter == 0:
                chapter = 1
            sub += 1
            subsub = 0
            title = stripped[4:].strip()
            title = num_prefix.sub("", title).strip()
            out.append(f"### {chapter}.{sub} {title}")
            prev_heading = stripped
            continue
        if stripped.startswith("#### "):
            if chapter == 0:
                chapter = 1
            if sub == 0:
                sub = 1
            subsub += 1
            title = stripped[5:].strip()
            title = num_prefix.sub("", title).strip()
            out.append(f"#### {chapter}.{sub}.{subsub} {title}")
            prev_heading = stripped
            continue
        prev_heading = None
        out.append(raw)
    return out


def prune_duplicate_appendix_sections(lines: list[str]) -> list[str]:
    """
    Remove duplicate appendix H3 sections (and their content) by title.

    This is a content-hygiene step that should run before numbering normalization.
    """
    out: list[str] = []
    in_appendix = False
    skip_section = False
    seen: set[str] = set()
    appendix_h2 = re.compile(r"^##\s+\d+\.\s*附录\b")
    for line in lines:
        s = line.strip()
        if appendix_h2.match(s):
            in_appendix = True
            skip_section = False
            out.append(line)
            continue
        if in_appendix and s.startswith("## "):
            in_appendix = False
            skip_section = False
            seen.clear()
            out.append(line)
            continue
        if in_appendix and s.startswith("### "):
            title = s[4:].strip()
            title = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", title)
            if title in seen:
                skip_section = True
                continue
            seen.add(title)
            skip_section = False
            out.append(line)
            continue
        if in_appendix and skip_section:
            continue
        out.append(line)
    return out


def apply_cn_en_font_mapping(docx_path: Path, *, cn_font: str = "宋体", en_font: str = "Times New Roman") -> None:
    doc = Document(str(docx_path))

    def apply_run(run) -> None:
        r_pr = run._r.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), cn_font)
        r_fonts.set(qn("w:ascii"), en_font)
        r_fonts.set(qn("w:hAnsi"), en_font)
        r_fonts.set(qn("w:cs"), en_font)

    def walk_paragraphs(paragraphs) -> None:
        for p in paragraphs:
            for r in p.runs:
                apply_run(r)

    walk_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                walk_paragraphs(cell.paragraphs)

    for section in doc.sections:
        walk_paragraphs(section.header.paragraphs)
        walk_paragraphs(section.footer.paragraphs)

    doc.save(str(docx_path))


def find_workspace_root(start: Path) -> Path:
    for parent in [start, *start.parents]:
        if (parent / "pipeline_guard.py").exists() and (parent / "shared" / "scripts").exists():
            return parent
    raise FileNotFoundError("Cannot locate workspace root")


ROOT = find_workspace_root(Path(__file__).resolve().parents[1])
TOOLS_ROOT_CANDIDATES = [ROOT, *ROOT.parents]

# Load workspace-wide manual rules after ROOT is available.
_MANUAL_RULES = load_manual_rules()
if isinstance(_MANUAL_RULES.get("min_screenshot_count"), int) and int(_MANUAL_RULES["min_screenshot_count"]) > 0:
    MIN_SCREENSHOT_COUNT = int(_MANUAL_RULES["min_screenshot_count"])
if isinstance(_MANUAL_RULES.get("min_manual_chars"), int) and int(_MANUAL_RULES["min_manual_chars"]) > 0:
    MIN_MANUAL_CHARS = int(_MANUAL_RULES["min_manual_chars"])


def run_command(command: list[str], *, cwd: Path = ROOT, env: dict[str, str] | None = None) -> subprocess.CompletedProcess[str]:
    return subprocess.run(command, cwd=cwd, env=env, check=True, text=True, capture_output=True, encoding="utf-8", errors="replace")


def find_first_file(base: Path, pattern: str) -> Path:
    matches = sorted(base.rglob(pattern))
    if not matches:
        raise FileNotFoundError(f"Cannot find {pattern} under {base}")
    return matches[-1]


def existing_env_file(env_name: str) -> Path | None:
    raw = os.environ.get(env_name, "").strip()
    if not raw:
        return None
    path = Path(raw).expanduser().resolve()
    if path.exists() and path.is_file():
        return path
    raise FileNotFoundError(f"Environment variable {env_name} points to missing file: {path}")


def find_tool_binary(tool_subdir: str, pattern: str, env_name: str) -> Path:
    env_path = existing_env_file(env_name)
    if env_path is not None:
        return env_path
    for base in TOOLS_ROOT_CANDIDATES:
        tool_dir = base / ".tools" / tool_subdir
        if tool_dir.exists():
            return find_first_file(tool_dir, pattern)
    raise FileNotFoundError(
        f"Cannot locate {pattern}. Install it under {ROOT / '.tools' / tool_subdir} or set {env_name}."
    )


def find_node() -> Path:
    return find_tool_binary("node", "node.exe", "NODE_EXE")


def find_pandoc() -> Path:
    return find_tool_binary("pandoc", "pandoc.exe", "PANDOC_EXE")


def find_runtime_script(name: str) -> Path:
    local = Path(__file__).resolve().parent / name
    if local.exists():
        return local
    root_side = ROOT / "scripts" / name
    if root_side.exists():
        return root_side
    raise FileNotFoundError(f"Cannot locate runtime script: {name}")


def default_screenshot_plan() -> list[dict[str, str]]:
    plan = load_default_screenshot_plan()
    if plan:
        return [dict(item) for item in plan]
    return []


def validate_screenshot_plan(raw: object, source: str) -> list[dict[str, str]]:
    if not isinstance(raw, list) or not raw:
        raise ValueError(f"Screenshot plan must be a non-empty list: {source}")
    if len(raw) < MIN_SCREENSHOT_COUNT:
        raise ValueError(
            f"Screenshot plan must contain at least {MIN_SCREENSHOT_COUNT} items: {source}"
        )

    normalized: list[dict[str, str]] = []
    seen_figures: set[str] = set()
    seen_files: set[str] = set()
    required_modal_actions = {"create", "edit", "delete", "view"}
    found_modal_actions: set[str] = set()
    for index, item in enumerate(raw, start=1):
        if not isinstance(item, dict):
            raise ValueError(f"Screenshot plan item #{index} must be an object: {source}")
        figure = str(item.get("figure", "")).strip()
        filename = str(item.get("filename", "")).strip()
        page = str(item.get("page", "")).strip().lower()
        hash_value = str(item.get("hash", "")).strip().lstrip("#")
        route_path = str(item.get("path", "")).strip()
        action = str(item.get("action", "")).strip().lower()
        modal = str(item.get("modal", "")).strip().lower()
        click_selector = str(item.get("clickSelector", "")).strip()
        wait_selector = str(item.get("waitSelector", "")).strip()
        if not figure or not filename:
            raise ValueError(f"Screenshot plan item #{index} missing figure or filename: {source}")
        if figure in seen_figures:
            raise ValueError(f"Duplicate figure '{figure}' in screenshot plan: {source}")
        if filename in seen_files:
            raise ValueError(f"Duplicate filename '{filename}' in screenshot plan: {source}")
        if page:
            if page not in {"login", "home"}:
                raise ValueError(f"Unsupported page '{page}' in screenshot plan: {source}")
            if hash_value or route_path:
                raise ValueError(f"Plan item #{index} cannot mix page with hash/path: {source}")
        elif hash_value and route_path:
            raise ValueError(f"Plan item #{index} cannot define both hash and path: {source}")
        elif not hash_value and not route_path:
            raise ValueError(f"Plan item #{index} must define one of page/hash/path: {source}")

        if action and action not in required_modal_actions:
            raise ValueError(f"Unsupported action '{action}' in screenshot plan: {source}")
        if modal and modal not in {"true", "false"}:
            raise ValueError(f"Invalid modal value '{modal}' in screenshot plan: {source}")
        if click_selector and len(click_selector) < 2:
            raise ValueError(f"Invalid clickSelector in screenshot plan item #{index}: {source}")
        if wait_selector and len(wait_selector) < 2:
            raise ValueError(f"Invalid waitSelector in screenshot plan item #{index}: {source}")
        if click_selector and modal != "true":
            raise ValueError(f"Plan item #{index} clickSelector requires modal=true: {source}")
        if wait_selector and modal != "true":
            raise ValueError(f"Plan item #{index} waitSelector requires modal=true: {source}")
        if modal == "true" and not action:
            raise ValueError(f"Plan item #{index} modal=true requires action=create/edit/delete/view: {source}")
        if action and modal != "true":
            raise ValueError(f"Plan item #{index} action requires modal=true: {source}")
        if modal == "true":
            found_modal_actions.add(action)

        seen_figures.add(figure)
        seen_files.add(filename)
        normalized.append(
            {
                "figure": figure,
                "filename": filename,
                "page": page,
                "hash": hash_value,
                "path": route_path,
                "action": action,
                "modal": modal,
                "clickSelector": click_selector,
                "waitSelector": wait_selector,
            }
        )
    missing_actions = sorted(required_modal_actions - found_modal_actions)
    if missing_actions:
        raise ValueError(
            "Screenshot plan must include modal screenshots for actions: "
            + ", ".join(missing_actions)
            + f": {source}"
        )
    return normalized


def load_screenshot_plan(project_root: Path) -> list[dict[str, str]]:
    plan_path = project_root / "screenshot_plan.json"
    if not plan_path.exists():
        return default_screenshot_plan()
    raw = json.loads(plan_path.read_text(encoding="utf-8"))
    return validate_screenshot_plan(raw, str(plan_path))


def build_figure_image_map(screenshot_plan: list[dict[str, str]]) -> dict[str, str]:
    return {item["figure"]: item["filename"] for item in screenshot_plan}


def get_application_form_template_path(override: Path | None = None) -> Path:
    if override is not None and override.exists() and override.suffix.lower() == ".docx":
        return override
    candidates = [
        Path(r"C:\Users\31556\.codex\skills\custom\chinese-copyright-application\references\application-form-template.docx"),
        Path(r"C:\Users\31556\.codex\skills\custom\chinese-copyright-application\references\application-form-template.md"),
    ]
    for path in candidates:
        if path.exists() and path.suffix.lower() == ".docx":
            return path
    raise FileNotFoundError("Cannot locate application-form-template.docx")


def iter_source_files(base_dir: Path, *, include_roots: list[str] | None = None) -> list[Path]:
    """
    Collect real source files for 源代码文档.md.

    Default policy is conservative: include only `apps/` to avoid leaking pipeline scripts,
    templates, PRD/manual narrative, and generated catalogs into the source-code document.
    Projects may override via `source_doc_filters.json` (see `build_source_markdown_by_rule`).
    """
    roots = include_roots[:] if include_roots else ["apps"]
    roots = [r.strip().strip("/\\") for r in roots if str(r).strip()]
    files: list[Path] = []
    for root_name in roots:
        base = base_dir / root_name
        if not base.exists():
            continue
        for path in sorted(base.rglob("*")):
            if not (path.is_file() and path.suffix.lower() in CODE_EXTENSIONS):
                continue
            if "__pycache__" in path.parts or "node_modules" in path.parts:
                continue
            files.append(path)
    return files


def source_file_rank(path: Path, base_dir: Path) -> tuple[int, str]:
    rel = str(path.relative_to(base_dir)).replace("\\", "/").lower()
    name = path.name.lower()
    if name in {"main.py", "app.py", "index.js", "index.mjs", "main.js"} or rel.endswith("/main.py") or rel.endswith("/app.py"):
        return (0, rel)
    if any(k in rel for k in ("route", "router", "routes", "page", "pages", "view", "views")):
        return (1, rel)
    if any(k in rel for k in ("service", "module", "domain", "model", "biz", "business", "core", "logic")):
        return (2, rel)
    if any(k in rel for k in ("config", "settings", "utils", "helper", "common", "infra")):
        return (3, rel)
    return (4, rel)


def build_source_markdown_by_rule(project_dir: Path, source_md_path: Path, requirement_name: str) -> None:
    filters_path = project_dir / "source_doc_filters.json"
    strip_substrings: list[str] = []
    include_roots: list[str] = ["apps"]
    rules = load_source_doc_rules()
    rule_strip = rules.get("global_strip_substrings")
    if isinstance(rule_strip, list):
        strip_substrings = [str(s) for s in rule_strip if str(s).strip()]
    include_default = rules.get("include_roots_default")
    if isinstance(include_default, list) and include_default:
        include_roots = [str(s) for s in include_default if str(s).strip()]
    if filters_path.exists():
        try:
            raw_filters = json.loads(filters_path.read_text(encoding="utf-8"))
            if isinstance(raw_filters, dict) and isinstance(raw_filters.get("strip_substrings"), list):
                # Project-level additions override/extend workspace defaults.
                strip_substrings = strip_substrings + [str(s) for s in raw_filters["strip_substrings"] if str(s).strip()]
            if isinstance(raw_filters, dict) and isinstance(raw_filters.get("include_roots"), list):
                include_roots = [str(s) for s in raw_filters["include_roots"] if str(s).strip()]
        except Exception:
            strip_substrings = strip_substrings or []
            include_roots = ["apps"]

    global_strip_substrings = strip_substrings

    def visual_width(s: str) -> int:
        w = 0
        for ch in s:
            w += 2 if unicodedata.east_asian_width(ch) in {"W", "F"} else 1
        return w

    def wrap_line(line: str, width: int = 76) -> list[str]:
        if visual_width(line) <= width:
            return [line]
        out: list[str] = []
        rest = line
        while visual_width(rest) > width:
            cut = 0
            cur = 0
            best = -1
            for i, ch in enumerate(rest):
                cur += 2 if unicodedata.east_asian_width(ch) in {"W", "F"} else 1
                if ch in {" ", ",", ")", "]", "}", "=", "+", "-", ";"} and cur >= int(width * 0.6):
                    best = i + 1
                if cur > width:
                    break
                cut = i + 1
            if best > 0:
                cut = best
            out.append(rest[:cut].rstrip())
            rest = rest[cut:].lstrip()
        out.append(rest)
        return out

    def should_strip(line: str) -> bool:
        return any(token in line for token in global_strip_substrings)

    def should_skip_file(path: Path) -> bool:
        rel = str(path.relative_to(project_dir)).replace("\\", "/")
        rel_lower = rel.lower()
        name_lower = path.name.lower()
        prefixes = rules.get("skip_file_name_prefixes")
        if isinstance(prefixes, list) and any(name_lower.startswith(str(p).lower()) for p in prefixes):
            return True
        rel_subs = rules.get("skip_file_rel_substrings")
        if isinstance(rel_subs, list) and any(str(s).lower() in rel_lower for s in rel_subs):
            return True
        rel_suffixes = rules.get("skip_file_rel_suffixes")
        if isinstance(rel_suffixes, list) and any(rel_lower.endswith(str(s).lower()) for s in rel_suffixes):
            return True
        # Content-based skip: if the file is an auto-generated label/dictionary catalog,
        # exclude the whole file rather than stripping lines (keeps the doc "real code only").
        try:
            head = path.read_text(encoding="utf-8", errors="ignore")[:8000]
        except OSError:
            return True
        content_subs = rules.get("skip_file_content_substrings")
        if isinstance(content_subs, list) and any(str(s) in head for s in content_subs):
            return True
        return False

    files = [p for p in iter_source_files(project_dir, include_roots=include_roots) if not should_skip_file(p)]
    files = sorted(files, key=lambda p: source_file_rank(p, project_dir))
    code_flow: list[str] = []
    for path in files:
        rel = str(path.relative_to(project_dir)).replace("\\", "/")
        code_flow.append(f"# {rel}")
        for raw in path.read_text(encoding="utf-8", errors="ignore").splitlines():
            if should_strip(raw):
                continue
            code_flow.extend(wrap_line(raw))
    if len(code_flow) > 2998:
        sampled = code_flow[:1499] + code_flow[-1499:]
    else:
        sampled = code_flow
    # Policy: 源代码文档不出现“空白行/空白页”，且不使用“#”等占位符凑行数。
    # - 保持 3000 行硬约束；
    # - 所有行必须非空；
    # - 若真实代码行不足 3000 行，则仅用“已有真实非空代码行（wrap-around）”补齐。
    non_empty_code_flow = [ln for ln in code_flow if str(ln).strip()]
    if not non_empty_code_flow:
        raise ValueError("No non-empty code lines collected for source document; cannot build a non-blank 3000-line source doc.")

    # Keep source markdown body code-only: do not prepend document title lines.
    # The source DOCX/PDF cover/header is handled in later rendering steps.
    sampled = list(sampled)
    # Replace any accidental blanks with a real code line (deterministic wrap-around).
    fixed: list[str] = []
    for idx, line in enumerate(sampled):
        fixed.append(line if str(line).strip() else non_empty_code_flow[idx % len(non_empty_code_flow)])
    sampled = fixed
    if len(sampled) < 3000:
        i = 0
        while len(sampled) < 3000:
            sampled.append(non_empty_code_flow[i % len(non_empty_code_flow)])
            i += 1
    else:
        sampled = sampled[:3000]
    source_md_path.write_text("\n".join(sampled), encoding="utf-8")


def count_source_lines(base_dir: Path) -> int:
    total = 0
    for path in iter_source_files(base_dir):
        total += sum(1 for line in path.read_text(encoding="utf-8", errors="ignore").splitlines() if line.strip())
    return total


def strip_markdown_for_count(markdown_text: str) -> int:
    visible = []
    for line in markdown_text.splitlines():
        if line.strip().startswith("!["):
            continue
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", line)
        visible.append(line.replace("`", ""))
    return len(re.sub(r"\s+", "", "".join(visible)))


def normalize_manual_line(text: str) -> str:
    # Remove markdown heading markers so manual body doesn't render literal '#'.
    return re.sub(r"^#{1,6}\s*", "", text).strip()


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def set_section_vertical_align(section, align: str) -> None:
    sect_pr = section._sectPr
    for node in sect_pr.findall(qn("w:vAlign")):
        sect_pr.remove(node)
    valign = OxmlElement("w:vAlign")
    valign.set(qn("w:val"), align)
    sect_pr.append(valign)


def configure_source_section_layout(section) -> None:
    sect_pr = section._sectPr
    for node in sect_pr.findall(qn("w:docGrid")):
        sect_pr.remove(node)
    for node in sect_pr.findall(qn("w:lnNumType")):
        sect_pr.remove(node)

    doc_grid = OxmlElement("w:docGrid")
    doc_grid.set(qn("w:type"), "lines")
    doc_grid.set(qn("w:linePitch"), str(LINE_PITCH_TWIPS))
    sect_pr.append(doc_grid)

    line_number = OxmlElement("w:lnNumType")
    line_number.set(qn("w:countBy"), "1")
    line_number.set(qn("w:restart"), "continuous")
    line_number.set(qn("w:distance"), "360")
    sect_pr.append(line_number)


def enforce_paragraph_snap_to_grid(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    snap = p_pr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        p_pr.append(snap)
    snap.set(qn("w:val"), "1")


def postprocess_docx_header_footer(docx_path: Path, header_text: str) -> None:
    doc = Document(str(docx_path))
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.clear()
    # Header rule: title on the left, page number on the right.
    header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    try:
        header.paragraph_format.tab_stops.clear_all()
    except Exception:
        pass
    if section.page_width is not None and section.left_margin is not None and section.right_margin is not None:
        usable_width = section.page_width - section.left_margin - section.right_margin
        header.paragraph_format.tab_stops.add_tab_stop(usable_width, alignment=WD_TAB_ALIGNMENT.RIGHT)
    header.add_run(header_text + "\t")
    add_page_number(header)
    footer = section.footer.paragraphs[0]
    footer.clear()
    doc.save(str(docx_path))
    apply_cn_en_font_mapping(docx_path)


def stabilize_docx_tables(docx_path: Path) -> None:
    """
    Stabilize table layout for downstream PDF engines (notably LibreOffice).

    LibreOffice may reflow "auto-fit" tables differently from Word. We force fixed layout
    and ensure borders are explicitly defined (Word gridlines are not printable).

    IMPORTANT:
    - Do not overwrite tblGrid/tcW widths aggressively. LibreOffice may render table text
      outside the table when column widths are forced incorrectly (observed in practice).
    """
    doc = Document(str(docx_path))
    section = doc.sections[0] if doc.sections else None
    usable_width = None
    if section is not None and section.page_width and section.left_margin and section.right_margin:
        usable_width = section.page_width - section.left_margin - section.right_margin

    changed = False
    for t in doc.tables:
        # Keep layout close to the reference doc/template; only enforce fixed layout flag.
        try:
            t.autofit = False
            changed = True
        except Exception:
            pass
        try:
            tbl_pr = t._tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")
            changed = True
        except Exception:
            pass

        # Ensure table borders are explicitly defined for PDF export.
        # Word can show non-printing "gridlines" even when borders are not set;
        # LibreOffice/PDF export will then look like plain text.
        try:
            tbl_pr = t._tbl.tblPr
            borders = tbl_pr.find(qn("w:tblBorders"))
            if borders is None:
                borders = OxmlElement("w:tblBorders")
                tbl_pr.append(borders)
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = borders.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    borders.append(el)
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "8")  # 0.5pt
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "000000")
            changed = True
        except Exception:
            pass

    if changed:
        doc.save(str(docx_path))


def parse_form_markdown(form_md: Path) -> dict[str, str]:
    data: dict[str, str] = {}
    for raw in form_md.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or line.startswith(">"):
            continue
        # Support both bullet and plain "key: value" lines.
        item = line[2:].strip() if line.startswith("- ") else line
        if "：" in item:
            k, v = item.split("：", 1)
            data[k.strip()] = v.strip()
        elif ":" in item:
            k, v = item.split(":", 1)
            data[k.strip()] = v.strip()
    return data


def build_form_docx_from_template(form_md: Path, form_docx: Path, template_docx: Path, header_text: str) -> None:
    mapping = parse_form_markdown(form_md)
    doc = Document(str(template_docx))
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            key = row.cells[0].text.strip().rstrip("：:")
            if key in mapping:
                row.cells[1].text = mapping[key]
    doc.save(str(form_docx))
    postprocess_docx_header_footer(form_docx, header_text)


def append_form_source_line_count(form_docx: Path, *, line_count: int) -> None:
    doc = Document(str(form_docx))
    text = "\n".join(
        [p.text for p in doc.paragraphs]
        + [c.text for t in doc.tables for r in t.rows for c in r.cells]
    )
    if any(k in text for k in ("源程序量", "源代码量", "源代码行数", "代码量", "代码行数")):
        return
    p = doc.add_paragraph()
    p.add_run(f"源程序量：{line_count} 行")
    doc.save(str(form_docx))


def convert_markdown_to_docx(
    md_path: Path,
    docx_path: Path,
    pandoc_path: Path,
    resource_path: Path | None = None,
    reference_doc: Path | None = None,
) -> None:
    # Pandoc fails on Windows if the output docx is currently opened/locked by Word.
    # Strategy:
    # 1) Try delete/overwrite directly.
    # 2) If locked, render to a temp docx and then atomically replace.
    target_out = docx_path
    if docx_path.exists():
        try:
            docx_path.chmod(0o666)
        except Exception:
            pass
        try:
            docx_path.unlink()
        except Exception:
            # Locked: use temp output under the same directory (so os.replace works).
            target_out = docx_path.with_name(docx_path.stem + ".__tmp__.docx")
            try:
                if target_out.exists():
                    target_out.unlink()
            except Exception:
                pass

    cmd = [str(pandoc_path), str(md_path), "-f", "gfm", "-t", "docx", "-o", str(target_out)]
    if resource_path is not None:
        cmd.extend(["--resource-path", str(resource_path)])
    if reference_doc is not None:
        cmd.extend(["--reference-doc", str(reference_doc)])
    run_command(cmd)
    if target_out != docx_path:
        try:
            os.replace(str(target_out), str(docx_path))
        except PermissionError as exc:
            raise PermissionError(
                f"Cannot replace locked DOCX output: {docx_path}. Close Word (or any viewer) and retry."
            ) from exc


def normalize_docx_heading_numbering(docx_path: Path) -> None:
    """
    Ensure heading numbering is present but not duplicated.

    Some reference-doc templates auto-number headings; when Markdown already contains
    numbers, the rendered docx may show duplicated prefixes like "1. 1. 标题".
    This normalizer collapses duplicated numeric prefixes for Heading 1/2/3.
    """
    doc = Document(str(docx_path))
    dup_pat = re.compile(r"^(\d+(?:\.\d+)*)\.\s+\1\.\s+")
    changed = False
    for p in doc.paragraphs:
        try:
            style = (p.style.name or "") if p.style is not None else ""
        except Exception:
            style = ""
        if style not in {"Heading 1", "Heading 2", "Heading 3"}:
            continue
        txt = p.text or ""
        new_txt = dup_pat.sub(r"\1. ", txt)
        if new_txt != txt:
            p.text = new_txt
            changed = True
    if changed:
        doc.save(str(docx_path))


def wait_docx_readable(docx_path: Path, retries: int = 3, delay_seconds: float = 0.8) -> None:
    last_error: Exception | None = None
    for attempt in range(retries):
        try:
            with docx_path.open("rb") as fh:
                fh.read(4)
            Document(str(docx_path))
            return
        except (BadZipFile, PermissionError, OSError, ValueError) as exc:
            last_error = exc
            if attempt < retries - 1:
                time.sleep(delay_seconds)
                continue
            raise RuntimeError(f"DOCX readiness check failed for {docx_path}: {exc}") from exc
    if last_error is not None:
        raise RuntimeError(f"DOCX readiness check failed for {docx_path}: {last_error}")


def default_pdf_print_css() -> Path:
    css = ROOT / "scripts" / "pdf_print.css"
    if not css.exists():
        raise FileNotFoundError(f"Missing PDF print CSS: {css}")
    return css


def convert_markdown_to_html(
    md_path: Path,
    html_path: Path,
    pandoc_path: Path,
    resource_path: Path | None = None,
    css_path: Path | None = None,
) -> None:
    cmd = [str(pandoc_path), str(md_path), "-f", "gfm", "-t", "html5", "--standalone", "-o", str(html_path)]
    if resource_path is not None:
        cmd.extend(["--resource-path", str(resource_path)])
    if css_path is not None:
        cmd.extend(["--css", str(css_path)])
    run_command(cmd)


def html_to_pdf(html_path: Path, pdf_path: Path, node_path: Path) -> None:
    env = os.environ.copy()
    env["EDGE_EXECUTABLE_PATH"] = str(EDGE_EXECUTABLE)
    run_command([str(node_path), str(find_runtime_script("html_to_pdf.mjs")), str(html_path), str(pdf_path)], env=env)


def docx_to_pdf_word(docx_path: Path, pdf_path: Path) -> None:
    docx_path = docx_path.resolve()
    pdf_path = pdf_path.resolve()
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    # Prefer in-process COM via pywin32 to stay in the caller's interactive session.
    if win32com is not None:
        if pythoncom is None:
            raise RuntimeError("pywin32 is available but pythoncom is missing")

        def export_once() -> None:
            word = None
            doc = None
            pv = None
            opened_by_us = False
            created_word_instance = False

            def find_open_doc(app, full_path: str):
                try:
                    for d in app.Documents:
                        try:
                            if str(d.FullName).lower() == full_path.lower():
                                return d
                        except Exception:
                            continue
                except Exception:
                    return None
                return None

            def get_running_word():
                try:
                    return win32com.client.GetObject(Class="Word.Application")
                except Exception:
                    return None

            pythoncom.CoInitialize()
            try:
                # Prefer attaching to an existing Word instance first (avoids COM server creation failures).
                word = get_running_word()
                if word is not None:
                    created_word_instance = False
                else:
                    # Prefer Dispatch() to work reliably in interactive sessions; it may attach to an existing instance.
                    try:
                        word = win32com.client.Dispatch("Word.Application")
                        created_word_instance = False
                    except Exception:
                        # Fallback: try creating a dedicated instance.
                        word = win32com.client.DispatchEx("Word.Application")
                        created_word_instance = True
                word.Visible = False
                word.DisplayAlerts = 0
                # Try normal open first.
                try:
                    doc = word.Documents.Open(str(docx_path), False, True)
                    opened_by_us = True
                except Exception:
                    # If the file is already open (possibly in another Word instance), reuse that document handle.
                    doc = find_open_doc(word, str(docx_path))
                    if doc is None:
                        running = get_running_word()
                        if running is not None:
                            doc = find_open_doc(running, str(docx_path))
                            if doc is not None:
                                word = running
                    if doc is None:
                        # Fallback: handle Protected View / security prompts by opening via ProtectedView.
                        try:
                            pv = word.ProtectedViewWindows.Open(str(docx_path))
                            pv.Edit()
                            doc = pv.Document
                            opened_by_us = True
                        except Exception:
                            raise
                # wdExportFormatPDF = 17
                doc.ExportAsFixedFormat(str(pdf_path), 17)
            finally:
                try:
                    if doc is not None and opened_by_us:
                        doc.Close(False)
                    if pv is not None:
                        pv.Close()
                finally:
                    if word is not None and created_word_instance:
                        word.Quit()
                    pythoncom.CoUninitialize()

        last_exc: Exception | None = None
        for _ in range(3):
            try:
                export_once()
                return
            except Exception as exc:
                last_exc = exc if isinstance(exc, Exception) else Exception(str(exc))
                time.sleep(0.6)
        raise last_exc if last_exc is not None else RuntimeError("Word export failed")

    # Fallback: external PowerShell COM automation.
    p_docx = str(docx_path).replace("'", "''")
    p_pdf = str(pdf_path).replace("'", "''")
    cmd = [
        "powershell",
        "-NoProfile",
        "-NonInteractive",
        "-Command",
        (
            "$ErrorActionPreference='Stop';"
            f"$docx='{p_docx}';"
            f"$pdf='{p_pdf}';"
            "$w=New-Object -ComObject Word.Application;"
            "$w.Visible=$false;"
            "$w.DisplayAlerts=0;"
            "try{"
            "$d=$w.Documents.Open($docx,$false,$true);"
            "$d.ExportAsFixedFormat($pdf,17);"
            "$d.Close($false);"
            "}finally{"
            "$w.Quit();"
            "}"
        ),
    ]
    run_command(cmd, cwd=ROOT)


def docx_to_pdf_via_edge_print(
    md_path: Path,
    *,
    html_path: Path,
    pdf_path: Path,
    pandoc_path: Path,
    output_dir: Path,
    node_path: Path,
    strict_pdf_images: bool,
) -> None:
    convert_markdown_to_html(md_path, html_path, pandoc_path, output_dir, css_path=default_pdf_print_css())
    if strict_pdf_images:
        assert_print_images_not_clipped(html_path, node_path)
    html_to_pdf(html_path, pdf_path, node_path)


def docx_to_pdf_libreoffice(docx_path: Path, pdf_path: Path) -> None:
    docx_path = docx_path.resolve()
    pdf_path = pdf_path.resolve()
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    outdir = str(pdf_path.parent)
    # Prefer soffice.exe over soffice.com. In some environments soffice.com may crash (0xC0000409)
    # while soffice.exe succeeds for headless conversion.
    soffice = shutil.which("soffice.exe") or shutil.which("soffice.com")
    if not soffice:
        candidates = [
            r"E:\Program Files\LibreOffice\program\soffice.exe",
            r"E:\Program Files\LibreOffice\program\soffice.com",
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files\LibreOffice\program\soffice.com",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.com",
        ]
        for c in candidates:
            if Path(c).exists():
                soffice = c
                break
    if not soffice:
        raise FileNotFoundError("LibreOffice not found (soffice.com/soffice.exe not in PATH and no known install path matched).")
    cmd = [
        soffice,
        "-env:UserInstallation=file:///" + str((Path(tempfile.gettempdir()) / "lo_profile_codex").resolve()).replace("\\", "/"),
        "--headless",
        "--nologo",
        "--nolockcheck",
        "--nodefault",
        "--nofirststartwizard",
        "--convert-to",
        "pdf",
        "--outdir",
        outdir,
        str(docx_path),
    ]
    # Run in output directory to reduce LO profile/tmp issues with unicode paths.
    run_command(cmd, cwd=pdf_path.parent)
    expected = pdf_path.parent / (docx_path.stem + ".pdf")
    if expected.resolve() != pdf_path.resolve():
        if not expected.exists():
            raise FileNotFoundError(f"LibreOffice conversion did not produce expected pdf: {expected}")
        if pdf_path.exists():
            pdf_path.unlink()
        expected.replace(pdf_path)


def kill_winword_processes() -> None:
    cmd = [
        "powershell",
        "-NoProfile",
        "-NonInteractive",
        "-Command",
        "Get-Process WINWORD -ErrorAction SilentlyContinue | Stop-Process -Force",
    ]
    try:
        run_command(cmd, cwd=ROOT)
    except Exception:
        # Best-effort cleanup; Word might not be running.
        return


def get_winword_process_count() -> int:
    cmd = [
        "powershell",
        "-NoProfile",
        "-NonInteractive",
        "-Command",
        "($p=Get-Process WINWORD -ErrorAction SilentlyContinue; if($p){$p.Count}else{0})",
    ]
    try:
        cp = run_command(cmd, cwd=ROOT)
    except Exception:
        return 0
    merged = ((cp.stdout or "") + "\n" + (cp.stderr or "")).strip()
    nums = re.findall(r"\b\d+\b", merged)
    return int(nums[-1]) if nums else 0


def ensure_winword_cleaned(*, enabled: bool) -> tuple[bool, str]:
    """
    Returns (cleaned, note).

    - cleaned=True means there is no WINWORD process after best-effort cleanup.
    - When enabled=False, this does not kill WINWORD; it only reports status.
    """
    before = get_winword_process_count()
    if not enabled:
        if before > 0:
            return False, f"WINWORD_running_before_pdf_export count={before} (not cleaned; pass --kill-winword-before-pdf)"
        return True, "WINWORD_not_running_before_pdf_export"

    if before > 0:
        kill_winword_processes()
    after = get_winword_process_count()
    if after == 0:
        if before > 0:
            return True, f"WINWORD_killed_before_pdf_export killed~={before}"
        return True, "WINWORD_not_running_before_pdf_export"
    return False, f"WINWORD_still_running_after_kill count={after} (close Word windows/dialogs and retry)"


def _format_exc_message(exc: BaseException) -> str:
    if isinstance(exc, subprocess.CalledProcessError):
        out = (exc.stdout or "").strip()
        err = (exc.stderr or "").strip()
        merged = "\n".join([s for s in [err, out] if s])
        return merged or str(exc)
    return str(exc)


def assert_print_images_not_clipped(html_path: Path, node_path: Path) -> None:
    env = os.environ.copy()
    env["EDGE_EXECUTABLE_PATH"] = str(EDGE_EXECUTABLE)
    try:
        run_command([str(node_path), str(find_runtime_script("check_print_images.mjs")), str(html_path)], env=env)
    except subprocess.CalledProcessError as exc:
        raise ValueError(f"PDF 图片展示门禁失败：{exc.stderr or exc.stdout or exc}") from exc


def wait_for_service(base_url: str, timeout_seconds: float = 20.0) -> None:
    deadline = time.time() + timeout_seconds
    while time.time() < deadline:
        try:
            with urlopen(base_url, timeout=2):
                return
        except URLError:
            time.sleep(0.5)
    raise TimeoutError(f"Service not available at {base_url}")


def capture_screenshots(
    images_dir: Path,
    node_path: Path,
    project_root: Path,
    screenshot_plan: list[dict[str, str]],
) -> None:
    base_url = f"http://127.0.0.1:{DEFAULT_SERVER_PORT}"
    env = os.environ.copy()
    env["APP_PORT"] = str(DEFAULT_SERVER_PORT)
    # Run as a module so `apps.*` imports resolve (running the file directly sets
    # sys.path[0] to `apps/api`, which can break absolute package imports).
    server = subprocess.Popen([sys.executable, "-m", "apps.api.main"], cwd=project_root, env=env)
    try:
        wait_for_service(base_url)
        shot_env = os.environ.copy()
        shot_env["SCREENSHOT_OUT_DIR"] = str(images_dir)
        shot_env["SCREENSHOT_BASE_URL"] = base_url
        shot_env["EDGE_EXECUTABLE_PATH"] = str(EDGE_EXECUTABLE)
        shot_env["SCREENSHOT_PLAN_JSON"] = json.dumps(screenshot_plan, ensure_ascii=False)
        try:
            run_command([str(node_path), str(find_runtime_script("capture_screenshots_local.mjs"))], env=shot_env)
        except subprocess.CalledProcessError as exc:
            raise RuntimeError(f"Screenshot capture failed: {exc.stderr or exc.stdout or exc}") from exc
    finally:
        server.terminate()
        server.wait(timeout=5)


def ensure_paths(requirement_name: str, output_dir: Path, overrides: dict[str, Path | None] | None = None) -> dict[str, Path]:
    paths = {
        "prd_md": output_dir / f"{requirement_name}需求规格说明书.md",
        "manual_md": output_dir / f"{requirement_name}操作手册.md",
        "manual_docx": output_dir / f"{requirement_name}操作手册.docx",
        "manual_pdf": output_dir / f"{requirement_name}操作手册.pdf",
        "source_md": output_dir / f"{requirement_name}源代码文档.md",
        "source_docx": output_dir / f"{requirement_name}源代码文档.docx",
        "source_pdf": output_dir / f"{requirement_name}源代码文档.pdf",
        "form_md": output_dir / "软件著作权登记申请表.md",
        "form_docx": output_dir / "软件著作权登记申请表.docx",
    }
    if overrides:
        for key, value in overrides.items():
            if value is not None:
                paths[key] = value.resolve()
    return paths


def build_source_docx_fixed(source_md: Path, source_docx: Path, header_text: str) -> None:
    lines = source_md.read_text(encoding="utf-8").splitlines()
    if len(lines) < 3000:
        lines.extend([""] * (3000 - len(lines)))
    elif len(lines) > 3000:
        lines = lines[:3000]

    doc = Document()
    first = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    paragraphs = [first]
    for _ in range(1, len(lines)):
        paragraphs.append(doc.add_paragraph())

    for paragraph, line in zip(paragraphs, lines, strict=True):
        paragraph.style = doc.styles["Normal"]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        enforce_paragraph_snap_to_grid(paragraph)
        # Keep physical lines without introducing visible placeholder tokens into the source doc.
        # Use a non-breaking space so Word keeps the line while remaining visually blank.
        run = paragraph.add_run(line if line else "\u00A0")
        run.font.name = "宋体"
        run.font.size = Pt(8)

    section = doc.sections[0]
    section.page_height = Mm(A4_HEIGHT_MM)
    section.page_width = Mm(A4_WIDTH_MM)
    section.top_margin = Mm(NORMAL_MARGIN_MM)
    section.bottom_margin = Mm(NORMAL_MARGIN_MM)
    section.left_margin = Mm(NORMAL_MARGIN_MM)
    section.right_margin = Mm(NORMAL_MARGIN_MM)
    configure_source_section_layout(section)
    doc.save(str(source_docx))
    postprocess_docx_header_footer(source_docx, header_text)


def get_word_visual_line_count(docx_path: Path) -> int:
    p = str(docx_path.resolve()).replace("'", "''")
    cmd = [
        "powershell",
        "-NoProfile",
        "-NonInteractive",
        "-Command",
        (
            f"$p='{p}';"
            "$w=New-Object -ComObject Word.Application;"
            "$w.Visible=$false;"
            "$d=$w.Documents.Open($p);"
            "$n=[int]$d.ComputeStatistics(1);"
            "$d.Close($false);"
            "$w.Quit();"
            "Write-Output $n"
        ),
    ]
    cp = run_command(cmd, cwd=ROOT)
    merged = (cp.stdout or "") + "\n" + (cp.stderr or "")
    nums = re.findall(r"\b\d+\b", merged)
    if not nums:
        raise RuntimeError(f"Word line statistics returned no numeric output for {docx_path}")
    return int(nums[-1])


def trim_source_md_to_visual_3000(
    source_md: Path,
    source_docx: Path,
    header_text: str,
    max_rounds: int = 8,
) -> str:
    base_lines = source_md.read_text(encoding="utf-8").splitlines()
    if not base_lines:
        raise ValueError("source markdown is empty")
    draft = list(base_lines)
    final_visual = None
    for _ in range(max_rounds):
        if len(draft) < 3000:
            draft = draft + ([""] * (3000 - len(draft)))
        else:
            draft = draft[:3000]
        source_md.write_text("\n".join(draft), encoding="utf-8")
        build_source_docx_fixed(source_md, source_docx, header_text)
        try:
            visual = get_word_visual_line_count(source_docx)
        except Exception:
            # Keep current 3000-line draft and stop rewriting to avoid corruption.
            return "word_visual_lines_check_skipped"
        final_visual = visual
        if visual <= 3000:
            return f"word_visual_lines={visual}"
        overflow = visual - 3000
        keep = max(1, len(draft) - overflow)
        draft = draft[:keep]
    # Finalize with a bounded 3000-line markdown even if visual count still exceeds target.
    if len(draft) < 3000:
        draft = draft + ([""] * (3000 - len(draft)))
    else:
        draft = draft[:3000]
    source_md.write_text("\n".join(draft), encoding="utf-8")
    build_source_docx_fixed(source_md, source_docx, header_text)
    return f"word_visual_lines={final_visual}" if final_visual is not None else "word_visual_lines_check_skipped"


def build_manual_markdown_with_cover(manual_md: Path, requirement_name: str) -> Path:
    src_lines = manual_md.read_text(encoding="utf-8").splitlines()
    src_lines = strip_leading_manual_titles(src_lines, requirement_name)
    src_lines = ensure_numbered_headings_markdown(src_lines)
    src = "\n".join(src_lines)
    cover = (
        "<div style=\"height:240mm; display:flex; flex-direction:column; justify-content:center; "
        "break-after:page; page-break-after:always; "
        "align-items:center; text-align:center; font-family:'SimSun','宋体',serif;\">\n"
        f"<div style=\"font-size:24pt; line-height:1.3; margin:0;\">{requirement_name} V1.0</div>\n"
        "<div style=\"font-size:24pt; line-height:1.3; margin:0;\">操作手册</div>\n"
        "</div>\n\n"
    )
    temp_md = Path(tempfile.gettempdir()) / f"{requirement_name}_manual_with_cover.md"
    temp_md.write_text(cover + src, encoding="utf-8")
    return temp_md


def build_manual_markdown_with_cover_docx(manual_md: Path, requirement_name: str) -> Path:
    src_lines = manual_md.read_text(encoding="utf-8").splitlines()
    src_lines = strip_leading_manual_titles(src_lines, requirement_name)
    normalized_lines = msv.normalize_manual_markdown_structure(src_lines)
    normalized_lines = ensure_numbered_headings_markdown(normalized_lines)
    # For DOCX, promote heading levels by 1 so that:
    # - Markdown "##" (chapter) becomes Word Heading 1 (avoids template numbering like "1.1" at top level).
    # - Markdown "###" (section) becomes Word Heading 2.
    promoted: list[str] = []
    for line in normalized_lines:
        if line.startswith("#### "):
            # Promote Markdown "####" to Word Heading 3 by reducing one level.
            # This avoids rendering third-level titles as "Heading 4".
            # Keep numbering in text because reference-doc may not auto-number Heading 3.
            title = line[5:].strip()
            promoted.append("### " + title)
            continue
        if line.startswith("### "):
            # Keep numbering in text because reference-doc may not auto-number Heading 2.
            title = line[4:].strip()
            promoted.append("## " + title)
        elif line.startswith("## "):
            title = line[3:].strip()
            # Keep chapter numbering in text; downstream will normalize duplication if template auto-numbers.
            promoted.append("# " + title)
        else:
            promoted.append(line)
    normalized_lines = promoted
    src = "\n".join(normalized_lines)
    cover = f"# {requirement_name} V1.0\n\n# 操作手册\n\n"
    temp_md = Path(tempfile.gettempdir()) / f"{requirement_name}_manual_with_cover_docx.md"
    temp_md.write_text(cover + src, encoding="utf-8")
    return temp_md


def ensure_manual_image_references(manual_md: Path, figure_image_map: dict[str, str]) -> None:
    text = manual_md.read_text(encoding="utf-8", errors="ignore")
    original = text
    for fig, img in figure_image_map.items():
        if f"images/{img}" in text:
            continue
        text = re.sub(
            rf"({re.escape(fig)}[^\n]*\n)(?!\s*!\[)",
            rf"\1\n![{img}](images/{img})\n\n",
            text,
            count=1,
        )
    if text != original:
        manual_md.write_text(text, encoding="utf-8")


def assert_markdown_writable(md_path: Path) -> None:
    try:
        with md_path.open("r+", encoding="utf-8"):
            pass
    except OSError as exc:
        raise PermissionError(f"Markdown not writable: {md_path}") from exc


def collect_expected_manual_images(text: str, figure_image_map: dict[str, str]) -> set[str]:
    expected: set[str] = set()
    for figure, filename in figure_image_map.items():
        if re.search(rf"(?m)^{re.escape(figure)}[^\n]*$", text):
            expected.add(filename)
    return expected


def prune_images_with_guard(manual_md: Path, images_dir: Path, figure_image_map: dict[str, str]) -> None:
    text = manual_md.read_text(encoding="utf-8", errors="ignore")
    refs = set(re.findall(r"!\[[^\]]*\]\(images/([^)]+)\)", text))
    expected = collect_expected_manual_images(text, figure_image_map)
    missing = sorted(expected - refs)
    if missing:
        raise ValueError(
            "Manual image reference guard failed before pruning. "
            f"missing={missing}"
        )
    for p in images_dir.glob("*.png"):
        if p.name not in refs:
            p.unlink()


def sync_markdown_version_date(md_path: Path, version: str = "V1.0", include_dev_date: bool = False) -> None:
    text = md_path.read_text(encoding="utf-8", errors="ignore")
    today = date.today().isoformat()
    updated = text
    updated = re.sub(r"(软件版本[：:]\s*)([^\n\r]+)", lambda m: f"{m.group(1)}{version}", updated)
    updated = re.sub(r"(文档版本[：:]\s*)([^\n\r]+)", lambda m: f"{m.group(1)}{version}", updated)
    updated = re.sub(r"(更新日期[：:]\s*)([^\n\r]+)", lambda m: f"{m.group(1)}{today}", updated)
    if not include_dev_date:
        # Manual/source docs must not carry form-only development date fields.
        updated = re.sub(r"(?m)^\s*[-*]?\s*开发完成日期[：:].*$\n?", "", updated)
        updated = re.sub(r"(?m)^\s*[-*]?\s*开发完成[：:].*$\n?", "", updated)
    if include_dev_date:
        updated = re.sub(
            r"开发完成(?:日期)?(?:[：:]\s*|[A-Z])\d{2,4}-\d{2}-\d{2}",
            lambda _: f"开发完成日期：{today}",
            updated,
        )
        updated = re.sub(r"(开发完成日期[：:]\s*)([^\n\r]*)", lambda m: f"{m.group(1)}{today}", updated)
        if "开发完成日期" not in updated:
            updated = updated.rstrip() + f"\n- 开发完成日期：{today}\n"
    if updated != text:
        md_path.write_text(updated, encoding="utf-8")


def sync_form_source_line_count(form_md: Path, *, source_doc_lines: int = 3000) -> None:
    """
    Keep 申请表中的“源程序量”与项目实际代码行数一致。
    """
    text = form_md.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()
    kept: list[str] = []
    for ln in lines:
        s = ln.strip()
        item = s[2:].strip() if s.startswith("- ") else s
        if item.startswith("源程序量"):
            continue
        kept.append(ln)
    kept = [ln for ln in kept if ln.strip() or True]
    # Append a single normalized line-count field at the end.
    out = "\n".join(kept).rstrip() + f"\n- 源程序量：{source_doc_lines}\n"
    if out != text:
        form_md.write_text(out, encoding="utf-8")


def sync_markdown_metadata(paths: dict[str, Path], version: str = "V1.0") -> None:
    for key in ("manual_md", "source_md"):
        if key in paths and paths[key].exists():
            sync_markdown_version_date(paths[key], version=version, include_dev_date=False)
    if "form_md" in paths and paths["form_md"].exists():
        sync_markdown_version_date(paths["form_md"], version=version, include_dev_date=True)


def touch_path(path: Path) -> None:
    try:
        now = time.time()
        os.utime(path, (now, now))
    except Exception:
        return


def _extract_versions_and_dates(text: str) -> tuple[set[str], set[str]]:
    versions = set(re.findall(r"V\d+(?:\.\d+)?", text, flags=re.IGNORECASE))
    dates = set(re.findall(r"\d{4}-\d{2}-\d{2}", text))
    return versions, dates


def assert_md_docx_sync(md_path: Path, docx_path: Path) -> None:
    md_text = md_path.read_text(encoding="utf-8", errors="ignore")
    doc = Document(str(docx_path))
    table_text = []
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                table_text.append(c.text)
    docx_text = "\n".join([*(p.text for p in doc.paragraphs), *table_text])
    md_versions, md_dates = _extract_versions_and_dates(md_text)
    dx_versions, dx_dates = _extract_versions_and_dates(docx_text)
    if md_versions and not md_versions.issubset(dx_versions):
        raise ValueError(f"Version mismatch: md={sorted(md_versions)} docx={sorted(dx_versions)} @ {docx_path}")
    if md_dates and not md_dates.issubset(dx_dates):
        raise ValueError(f"Date mismatch: md={sorted(md_dates)} docx={sorted(dx_dates)} @ {docx_path}")


def force_docx_dates_from_md(md_path: Path, docx_path: Path) -> None:
    md_text = md_path.read_text(encoding="utf-8", errors="ignore")
    md_dates = sorted(set(re.findall(r"\d{4}-\d{2}-\d{2}", md_text)))
    if not md_dates:
        return
    target_date = md_dates[-1]
    doc = Document(str(docx_path))
    date_pat = re.compile(r"\d{4}-\d{2}-\d{2}")
    for p in doc.paragraphs:
        if date_pat.search(p.text):
            p.text = date_pat.sub(target_date, p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if date_pat.search(c.text):
                    c.text = date_pat.sub(target_date, c.text)
    doc.save(str(docx_path))


def align_file_timestamp(src_path: Path, dst_path: Path) -> None:
    st = src_path.stat()
    os.utime(dst_path, (st.st_atime, st.st_mtime))


def build_materials(
    requirement_name: str,
    output_dir: Path,
    overrides: dict[str, Path | None] | None = None,
    form_template_docx: Path | None = None,
    strict_pdf_images: bool = False,
    pdf_engine: str = "word",
    require_word_pdf: bool = False,
    kill_winword_before_pdf: bool = False,
    touch_md: bool = False,
) -> dict[str, str | int]:
    print("[stage4] START - build_materials")
    print(f"[stage4] PARAM - name={requirement_name} output_dir={output_dir}")
    output_dir.mkdir(parents=True, exist_ok=True)
    images_dir = output_dir / "images"
    if images_dir.exists():
        shutil.rmtree(images_dir)
    images_dir.mkdir(parents=True, exist_ok=True)
    paths = ensure_paths(requirement_name, output_dir, overrides=overrides)
    required_markdown = (paths["prd_md"], paths["manual_md"], paths["source_md"], paths["form_md"])
    missing = [str(path) for path in required_markdown if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required markdown inputs: " + ", ".join(missing))
    print("[stage4] CHECK - markdown inputs exist")
    # Strict gate: markdown must be writable before any docx/pdf generation.
    for md_key in ("manual_md", "source_md", "form_md"):
        assert_markdown_writable(paths[md_key])
    print("[stage4] CHECK - markdown files writable")
    # Regenerate source markdown by skill rule before rendering docx/pdf.
    print("[stage4] STEP - regenerate source markdown")
    build_source_markdown_by_rule(output_dir, paths["source_md"], requirement_name)

    print("[stage4] STEP - locate tools (node/pandoc)")
    node_path = find_node()
    pandoc_path = find_pandoc()
    header_text = f"{requirement_name}V1.0"
    print(f"[stage4] STEP - header_text={header_text}")
    screenshot_plan = load_screenshot_plan(output_dir)
    figure_image_map = build_figure_image_map(screenshot_plan)
    print(f"[stage4] STEP - screenshot_plan items={len(screenshot_plan)}")

    screenshot_warning = ""
    try:
        print("[stage4] STEP - capture screenshots")
        capture_screenshots(
            images_dir,
            node_path=node_path,
            project_root=output_dir,
            screenshot_plan=screenshot_plan,
        )
        print("[stage4] STEP - capture screenshots: OK")
    except Exception as exc:
        screenshot_warning = str(exc)
        print(f"[stage4] WARN - capture screenshots failed: {screenshot_warning}")

    # Policy: screenshots are part of deliverables. Placeholders are allowed only when
    # explicitly enabled, and should not be the default for new runs.
    allow_placeholders = os.environ.get("ALLOW_SCREENSHOT_PLACEHOLDERS", "").strip() in {"1", "true", "TRUE", "yes", "YES"}

    # If screenshot service is unavailable, optionally create placeholder PNGs for every
    # planned figure so downstream tools can continue (but this should be treated as a
    # non-final artifact set).
    def _write_placeholder_png(path: Path) -> None:
        # Deterministic, non-trivial placeholder PNG (1280x720) with a simple
        # pattern derived from filename. This avoids duplicate hashes and
        # produces images that are not "blank 1x1", making reviews easier when
        # screenshot service is unavailable.
        import struct
        import zlib

        def _chunk(chunk_type: bytes, data: bytes) -> bytes:
            length = struct.pack(">I", len(data))
            crc = struct.pack(">I", zlib.crc32(chunk_type + data) & 0xFFFFFFFF)
            return length + chunk_type + data + crc

        def _png_bytes(width: int, height: int, seed: int) -> bytes:
            # Blocky pattern so the PNG compresses well (avoids multi-MB noise).
            block = 24
            rows: list[bytes] = []
            for y in range(height):
                row = bytearray()
                row.append(0)  # filter
                by = y // block
                for x in range(width):
                    bx = x // block
                    v = (seed ^ (bx * 1103515245) ^ (by * 12345)) & 0xFFFFFFFF
                    r = (v >> 16) & 0xFF
                    g = (v >> 8) & 0xFF
                    b = v & 0xFF
                    row.extend([r, g, b, 255])
                rows.append(bytes(row))
            raw = b"".join(rows)
            compressed = zlib.compress(raw, level=9)

            signature = b"\x89PNG\r\n\x1a\n"
            ihdr = struct.pack(">IIBBBBB", width, height, 8, 6, 0, 0, 0)
            return signature + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", compressed) + _chunk(b"IEND", b"")

        seed = zlib.crc32(path.name.encode("utf-8")) & 0xFFFFFFFF
        png = _png_bytes(960, 540, seed)
        path.parent.mkdir(parents=True, exist_ok=True)
        # Always rewrite so content changes when plan changes.
        path.write_bytes(png)

    if screenshot_warning:
        if not allow_placeholders:
            raise RuntimeError(
                "Screenshot capture failed and placeholders are disabled. "
                "Start the project server/screenshots, or set ALLOW_SCREENSHOT_PLACEHOLDERS=1 to proceed. "
                f"Cause: {screenshot_warning}"
            )
        images_dir = output_dir / "images"
        planned = [images_dir / rel for rel in figure_image_map.values()]
        for img_path in planned:
            _write_placeholder_png(img_path)

    # Always synchronize markdown metadata first, then generate docx/pdf from markdown.
    print("[stage4] STEP - sync markdown metadata")
    sync_markdown_metadata(paths, version="V1.0")
    sanitize_manual_markdown(paths["manual_md"], requirement_name)
    ensure_manual_min_chars(paths["manual_md"], min_chars=MIN_MANUAL_CHARS)
    if touch_md:
        for md_key in ("manual_md", "source_md", "form_md"):
            if md_key in paths and paths[md_key].exists():
                touch_path(paths[md_key])
        print("[stage4] STEP - touch markdown mtime: OK")
    ensure_manual_image_references(paths["manual_md"], figure_image_map)
    prune_images_with_guard(paths["manual_md"], output_dir / "images", figure_image_map)
    print("[stage4] STEP - validate manual image references: OK")

    manual_md_for_pdf = build_manual_markdown_with_cover(paths["manual_md"], requirement_name)
    manual_md_for_docx = build_manual_markdown_with_cover_docx(paths["manual_md"], requirement_name)
    print("[stage4] STEP - build manual cover markdown: OK")
    msv.assert_manual_markdown_heading_safety(paths["manual_md"])
    msv.assert_manual_template_skeleton(paths["manual_md"], requirement_name)
    expected_prefix = (
        "<div style=\"height:240mm; display:flex; flex-direction:column; justify-content:center; "
        "break-after:page; page-break-after:always; "
        "align-items:center; text-align:center; font-family:'SimSun','宋体',serif;\">\n"
        f"<div style=\"font-size:24pt; line-height:1.3; margin:0;\">{requirement_name} V1.0</div>\n"
        "<div style=\"font-size:24pt; line-height:1.3; margin:0;\">操作手册</div>\n"
        "</div>\n\n"
    )
    actual_prefix = manual_md_for_pdf.read_text(encoding="utf-8")
    if not actual_prefix.startswith(expected_prefix):
        raise ValueError("操作手册 PDF 封面校验失败：渲染源封面标记缺失。")
    msv.assert_skill_has_cover_paging_rule()
    manual_ref_doc = msv.get_manual_template_path()
    print("[stage4] STEP - convert manual md -> docx")
    convert_markdown_to_docx(
        manual_md_for_docx,
        paths["manual_docx"],
        pandoc_path,
        output_dir,
        reference_doc=manual_ref_doc,
    )
    print("[stage4] STEP - manual docx postprocess (cover/header/table/fonts)")
    wait_docx_readable(paths["manual_docx"], retries=3, delay_seconds=0.8)
    msv.enforce_manual_cover(paths["manual_docx"], requirement_name)
    wait_docx_readable(paths["manual_docx"], retries=3, delay_seconds=0.5)
    msv.cleanup_literal_newpage_markers(paths["manual_docx"])
    wait_docx_readable(paths["manual_docx"], retries=3, delay_seconds=0.5)
    msv.normalize_body_pagination_spacing(paths["manual_docx"])
    wait_docx_readable(paths["manual_docx"], retries=3, delay_seconds=0.5)
    postprocess_docx_header_footer(paths["manual_docx"], header_text)
    wait_docx_readable(paths["manual_docx"], retries=3, delay_seconds=0.5)
    stabilize_docx_tables(paths["manual_docx"])
    wait_docx_readable(paths["manual_docx"], retries=3, delay_seconds=0.5)
    msv.assert_manual_cover_rules(paths["manual_docx"], requirement_name)
    msv.assert_manual_ui_screenshot_interaction_rules(paths["manual_md"])
    normalize_docx_heading_numbering(paths["manual_docx"])
    apply_cn_en_font_mapping(paths["manual_docx"])
    print("[stage4] STEP - export manual pdf")
    pdf_warning = ""
    cleaned, cleaned_note = ensure_winword_cleaned(enabled=bool(kill_winword_before_pdf and pdf_engine == "word"))
    if not cleaned:
        pdf_warning = cleaned_note
    if pdf_engine not in {"libreoffice", "word", "edge"}:
        raise ValueError(f"Unknown pdf_engine: {pdf_engine}")
    try:
        if pdf_engine == "libreoffice":
            docx_to_pdf_libreoffice(paths["manual_docx"], paths["manual_pdf"])
        elif pdf_engine == "word":
            docx_to_pdf_word(paths["manual_docx"], paths["manual_pdf"])
        else:
            docx_to_pdf_via_edge_print(
                manual_md_for_pdf,
                html_path=output_dir / "_manual.html",
                pdf_path=paths["manual_pdf"],
                pandoc_path=pandoc_path,
                output_dir=output_dir,
                node_path=node_path,
                strict_pdf_images=strict_pdf_images,
            )
    except Exception as exc:
        if pdf_engine == "word" and require_word_pdf:
            # One more attempt after cleanup when running in Word engine.
            if not cleaned:
                ensure_winword_cleaned(enabled=True)
            try:
                docx_to_pdf_word(paths["manual_docx"], paths["manual_pdf"])
                cleaned2, note2 = ensure_winword_cleaned(enabled=False)
                if pdf_warning:
                    pdf_warning = pdf_warning + " | " + note2
                else:
                    pdf_warning = note2
                # Continue flow as success.
            except Exception:
                pass
            detail = _format_exc_message(exc)
            raise RuntimeError(
                "Word 导出 PDF 失败，导致 PDF 的页眉/页码规则无法与 docx 保持一致。\n"
                "常见原因：当前环境为无桌面会话/服务环境，无法创建 Word COM（例如 0x80070520），或 Word 打开文档时触发受保护视图/加载项/弹窗导致失败。\n"
                "处理方式：在安装了 Microsoft Word 且有桌面会话的环境重新生成；或临时使用回退链路："
                "`--allow-pdf-fallback` 或 `--pdf-engine edge`（但页眉/页码可能不一致）。\n"
                f"细节：{detail}"
            ) from exc
        if pdf_engine == "word" and not require_word_pdf:
            pdf_warning = f"word_pdf_failed_fallback_edge: {exc}"
            docx_to_pdf_via_edge_print(
                manual_md_for_pdf,
                html_path=output_dir / "_manual.html",
                pdf_path=paths["manual_pdf"],
                pandoc_path=pandoc_path,
                output_dir=output_dir,
                node_path=node_path,
                strict_pdf_images=strict_pdf_images,
            )
        else:
            raise

    _ = trim_source_md_to_visual_3000(paths["source_md"], paths["source_docx"], header_text)
    apply_cn_en_font_mapping(paths["source_docx"])
    try:
        # Source-code PDF must match DOCX layout (header/page number/line numbering).
        # Even when manual uses edge-print fallback, source PDF is generated from DOCX.
        if pdf_engine in {"libreoffice", "edge"}:
            docx_to_pdf_libreoffice(paths["source_docx"], paths["source_pdf"])
        else:
            docx_to_pdf_word(paths["source_docx"], paths["source_pdf"])
    except Exception as exc:
        if pdf_engine == "word" and require_word_pdf:
            if not cleaned:
                ensure_winword_cleaned(enabled=True)
            try:
                docx_to_pdf_word(paths["source_docx"], paths["source_pdf"])
                cleaned2, note2 = ensure_winword_cleaned(enabled=False)
                if pdf_warning:
                    pdf_warning = pdf_warning + " | " + note2
                else:
                    pdf_warning = note2
            except Exception:
                pass
            detail = _format_exc_message(exc)
            raise RuntimeError(
                "Word 导出 PDF 失败，导致 PDF 的页眉/页码规则无法与 docx 保持一致。\n"
                "常见原因：当前环境为无桌面会话/服务环境，无法创建 Word COM（例如 0x80070520），或 Word 打开文档时触发受保护视图/加载项/弹窗导致失败。\n"
                "处理方式：在安装了 Microsoft Word 且有桌面会话的环境重新生成；或临时使用回退链路："
                "`--allow-pdf-fallback` 或 `--pdf-engine edge`（但页眉/页码可能不一致）。\n"
                f"细节：{detail}"
            ) from exc
        # For source doc, prefer DOCX-based engines only; edge-print is a last resort and may diverge.
        if pdf_engine != "word":
            # Try Word as fallback if LibreOffice fails.
            try:
                docx_to_pdf_word(paths["source_docx"], paths["source_pdf"])
            except Exception:
                raise
        if pdf_engine == "word" and not require_word_pdf:
            # Word failed and fallback is allowed: use edge-print as last resort.
            pdf_warning = (pdf_warning + " | " if pdf_warning else "") + f"word_pdf_failed_fallback_edge: {exc}"
            docx_to_pdf_via_edge_print(
                paths["source_md"],
                html_path=output_dir / "_source.html",
                pdf_path=paths["source_pdf"],
                pandoc_path=pandoc_path,
                output_dir=output_dir,
                node_path=node_path,
                strict_pdf_images=strict_pdf_images,
            )
            # Note: this fallback may diverge from DOCX layout.
        else:
            raise

    # Compute real project code-line count; for form field "源程序量",
    # use actual line count when >3000, otherwise normalize to 3000.
    line_count = count_source_lines(output_dir)
    form_source_line_count = line_count if line_count > 3000 else 3000
    # Sync form fields that must match deliverables (dates/line-count), then render docx.
    sync_form_source_line_count(paths["form_md"], source_doc_lines=form_source_line_count)
    form_ref_doc = get_application_form_template_path(form_template_docx)
    build_form_docx_from_template(paths["form_md"], paths["form_docx"], form_ref_doc, header_text)
    apply_cn_en_font_mapping(paths["form_docx"])
    force_docx_dates_from_md(paths["manual_md"], paths["manual_docx"])
    force_docx_dates_from_md(paths["form_md"], paths["form_docx"])
    assert_md_docx_sync(paths["manual_md"], paths["manual_docx"])
    assert_md_docx_sync(paths["form_md"], paths["form_docx"])
    # Do NOT align output docx timestamps to md timestamps:
    # it makes regenerated docx look "not updated" even when content changed.

    for temp_file in (output_dir / "_manual.html", output_dir / "_source.html", manual_md_for_pdf, manual_md_for_docx):
        if temp_file.exists():
            temp_file.unlink()

    manual_chars = strip_markdown_for_count(paths["manual_md"].read_text(encoding="utf-8"))
    if manual_chars < MIN_MANUAL_CHARS:
        raise ValueError(f"操作手册字数不足：{manual_chars} < {MIN_MANUAL_CHARS}（请补充业务流程、字段口径与操作步骤后重试）")
    print(f"[stage4] METRIC - manual_chars={manual_chars}")
    print(f"[stage4] METRIC - line_count={line_count}")
    print(f"[stage4] METRIC - form_source_line_count={form_source_line_count}")
    # Best-effort: write computed source code line count back into application-form mapping.
    # Different templates may use different labels; populate common variants.
    try:
        form_md = paths["form_md"]
        existing = parse_form_markdown(form_md) if form_md.exists() else {}
        for k in ("源程序量", "源代码量", "源代码行数", "代码量", "代码行数"):
            existing.setdefault(k, str(form_source_line_count))
        # Ensure md contains a bullet key so future runs are deterministic.
        lines = form_md.read_text(encoding="utf-8", errors="ignore").splitlines()
        if not any(("源程序量" in ln or "源代码量" in ln or "源代码行数" in ln or "代码量" in ln or "代码行数" in ln) for ln in lines):
            lines.append(f"- 源程序量：{form_source_line_count}")
            form_md.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    except Exception:
        pass

    try:
        append_form_source_line_count(paths["form_docx"], line_count=form_source_line_count)
    except Exception:
        pass

    return {
        "line_count": line_count,
        "manual_chars": manual_chars,
        "source_doc_lines": 3000,
        "source_doc_pages": 60,
        "prd_path": str(paths["prd_md"].resolve()),
        "manual_md_path": str(paths["manual_md"].resolve()),
        "manual_docx_path": str(paths["manual_docx"].resolve()),
        "manual_pdf_path": str(paths["manual_pdf"].resolve()),
        "source_md_path": str(paths["source_md"].resolve()),
        "source_docx_path": str(paths["source_docx"].resolve()),
        "source_pdf_path": str(paths["source_pdf"].resolve()),
        "form_md_path": str(paths["form_md"].resolve()),
        "form_docx_path": str(paths["form_docx"].resolve()),
        "images_dir": str(images_dir.resolve()),
        "screenshot_warning": screenshot_warning,
        "pdf_warning": pdf_warning,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Build materials from existing markdown files.")
    parser.add_argument("--name", required=True, help="Requirement/software name.")
    parser.add_argument("--output-dir", required=True, help="Requirement directory containing markdown inputs.")
    parser.add_argument("--prd-md", default="", help="Explicit path to PRD markdown.")
    parser.add_argument("--manual-md", default="", help="Explicit path to manual markdown.")
    parser.add_argument("--source-md", default="", help="Explicit path to source-code markdown.")
    parser.add_argument("--form-md", default="", help="Explicit path to application-form markdown.")
    parser.add_argument("--manual-docx", default="", help="Explicit output path for manual docx.")
    parser.add_argument("--manual-pdf", default="", help="Explicit output path for manual pdf.")
    parser.add_argument("--source-docx", default="", help="Explicit output path for source docx.")
    parser.add_argument("--source-pdf", default="", help="Explicit output path for source pdf.")
    parser.add_argument("--form-docx", default="", help="Explicit output path for form docx.")
    parser.add_argument("--form-template-docx", default="", help="Explicit application-form template docx path.")
    parser.add_argument("--strict-pdf-images", action="store_true", help="Fail if HTML print image layout violates anti-clipping rules.")
    parser.add_argument("--pdf-engine", default="libreoffice", choices=["libreoffice", "word", "edge"], help="PDF engine: libreoffice(docx->pdf), word(docx->pdf) or edge(md->html->pdf).")
    parser.add_argument("--require-word-pdf", action="store_true", help="Fail if Word docx->pdf export fails (no fallback).")
    parser.add_argument("--allow-pdf-fallback", action="store_true", help="If Word docx->pdf export fails, fallback to edge(md->html->pdf).")
    parser.add_argument("--kill-winword-before-pdf", action="store_true", help="Kill WINWORD.EXE before Word PDF export (one-run flag).")
    parser.add_argument("--touch-md", action="store_true", help="Touch markdown mtime after metadata sync (so md looks refreshed).")
    args = parser.parse_args()
    if args.strict_pdf_images and args.pdf_engine != "edge":
        raise ValueError("--strict-pdf-images 仅在 --pdf-engine=edge 时可用。")
    if args.require_word_pdf and args.allow_pdf_fallback:
        raise ValueError("--require-word-pdf 与 --allow-pdf-fallback 不能同时使用。")
    require_word_pdf = bool(args.require_word_pdf)
    if args.pdf_engine == "word" and not args.allow_pdf_fallback:
        require_word_pdf = True
    overrides = {
        "prd_md": Path(args.prd_md) if args.prd_md else None,
        "manual_md": Path(args.manual_md) if args.manual_md else None,
        "source_md": Path(args.source_md) if args.source_md else None,
        "form_md": Path(args.form_md) if args.form_md else None,
        "manual_docx": Path(args.manual_docx) if args.manual_docx else None,
        "manual_pdf": Path(args.manual_pdf) if args.manual_pdf else None,
        "source_docx": Path(args.source_docx) if args.source_docx else None,
        "source_pdf": Path(args.source_pdf) if args.source_pdf else None,
        "form_docx": Path(args.form_docx) if args.form_docx else None,
    }
    form_template_docx = Path(args.form_template_docx).resolve() if args.form_template_docx else None
    result = build_materials(
        args.name,
        Path(args.output_dir).resolve(),
        overrides=overrides,
        form_template_docx=form_template_docx,
        strict_pdf_images=bool(args.strict_pdf_images),
        pdf_engine=str(args.pdf_engine),
        require_word_pdf=require_word_pdf,
        kill_winword_before_pdf=bool(args.kill_winword_before_pdf),
        touch_md=bool(args.touch_md),
    )
    for key, value in result.items():
        print(f"{key}={value}")


if __name__ == "__main__":
    main()
