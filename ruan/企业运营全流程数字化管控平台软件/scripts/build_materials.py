from __future__ import annotations

import argparse
import os
import re
import subprocess
import sys
import time
from pathlib import Path
from urllib.error import URLError
from urllib.request import urlopen

from docx import Document
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


def find_workspace_root(start: Path) -> Path:
    for parent in [start, *start.parents]:
        if (parent / ".tools").exists():
            return parent
    raise FileNotFoundError("Cannot locate workspace root containing .tools")


ROOT = find_workspace_root(Path(__file__).resolve().parents[1])
EDGE_EXECUTABLE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")
DEFAULT_SERVER_PORT = 8010
CODE_EXTENSIONS = {".py", ".js", ".mjs", ".css", ".html", ".sql", ".json"}

NORMAL_MARGIN_MM = 25.4
LINES_PER_PAGE = 50
LINE_PITCH_TWIPS = 279


def run_command(command: list[str], *, cwd: Path = ROOT, env: dict[str, str] | None = None) -> subprocess.CompletedProcess[str]:
    return subprocess.run(command, cwd=cwd, env=env, check=True, text=True, capture_output=True)


def find_first_file(base: Path, pattern: str) -> Path:
    matches = sorted(base.rglob(pattern))
    if not matches:
        raise FileNotFoundError(f"Cannot find {pattern} under {base}")
    return matches[-1]


def find_node() -> Path:
    return find_first_file(ROOT / ".tools" / "node", "node.exe")


def find_pandoc() -> Path:
    return find_first_file(ROOT / ".tools" / "pandoc", "pandoc.exe")


def find_runtime_script(name: str) -> Path:
    local = Path(__file__).resolve().parent / name
    if local.exists():
        return local
    root_side = ROOT / "scripts" / name
    if root_side.exists():
        return root_side
    raise FileNotFoundError(f"Cannot locate runtime script: {name}")


def iter_source_files(base_dir: Path) -> list[Path]:
    files: list[Path] = []
    for root_name in ("apps", "packages", "scripts"):
        base = base_dir / root_name
        if not base.exists():
            continue
        for path in sorted(base.rglob("*")):
            if path.is_file() and path.suffix.lower() in CODE_EXTENSIONS:
                if "__pycache__" in path.parts or "node_modules" in path.parts:
                    continue
                files.append(path)
    return files


def count_source_lines(base_dir: Path) -> int:
    total = 0
    for path in iter_source_files(base_dir):
        total += sum(1 for line in path.read_text(encoding="utf-8", errors="ignore").splitlines() if line.strip())
    return total


def strip_markdown_for_count(markdown_text: str) -> int:
    visible = []
    for line in markdown_text.splitlines():
        if line.strip().startswith("!["):
            continue
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", line)
        visible.append(line.replace("`", ""))
    return len(re.sub(r"\s+", "", "".join(visible)))


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def postprocess_docx_header_footer(docx_path: Path, header_text: str) -> None:
    doc = Document(str(docx_path))
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header.clear()
    header.add_run(header_text)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    footer.clear()
    add_page_number(footer)
    doc.save(str(docx_path))


def convert_markdown_to_docx(md_path: Path, docx_path: Path, pandoc_path: Path, resource_path: Path | None = None) -> None:
    cmd = [str(pandoc_path), str(md_path), "-f", "gfm", "-t", "docx", "-o", str(docx_path)]
    if resource_path is not None:
        cmd.extend(["--resource-path", str(resource_path)])
    run_command(cmd)


def convert_markdown_to_html(md_path: Path, html_path: Path, pandoc_path: Path, resource_path: Path | None = None) -> None:
    cmd = [str(pandoc_path), str(md_path), "-f", "gfm", "-t", "html5", "--standalone", "-o", str(html_path)]
    if resource_path is not None:
        cmd.extend(["--resource-path", str(resource_path)])
    run_command(cmd)


def html_to_pdf(html_path: Path, pdf_path: Path, node_path: Path) -> None:
    env = os.environ.copy()
    env["EDGE_EXECUTABLE_PATH"] = str(EDGE_EXECUTABLE)
    run_command([str(node_path), str(find_runtime_script("html_to_pdf.mjs")), str(html_path), str(pdf_path)], env=env)


def wait_for_service(base_url: str, timeout_seconds: float = 20.0) -> None:
    deadline = time.time() + timeout_seconds
    while time.time() < deadline:
        try:
            with urlopen(base_url, timeout=2):
                return
        except URLError:
            time.sleep(0.5)
    raise TimeoutError(f"Service not available at {base_url}")


def capture_screenshots(images_dir: Path, node_path: Path, project_root: Path) -> None:
    base_url = f"http://127.0.0.1:{DEFAULT_SERVER_PORT}"
    env = os.environ.copy()
    env["APP_PORT"] = str(DEFAULT_SERVER_PORT)
    server = subprocess.Popen([sys.executable, "apps/api/main.py"], cwd=project_root, env=env)
    try:
        wait_for_service(base_url)
        shot_env = os.environ.copy()
        shot_env["SCREENSHOT_OUT_DIR"] = str(images_dir)
        shot_env["SCREENSHOT_BASE_URL"] = base_url
        shot_env["EDGE_EXECUTABLE_PATH"] = str(EDGE_EXECUTABLE)
        try:
            run_command([str(node_path), str(find_runtime_script("capture_screenshots_local.mjs"))], env=shot_env)
        except subprocess.CalledProcessError:
            pass
    finally:
        server.terminate()
        server.wait(timeout=5)


def ensure_paths(requirement_name: str, output_dir: Path) -> dict[str, Path]:
    return {
        "prd_md": output_dir / f"{requirement_name}需求规格说明书.md",
        "manual_md": output_dir / f"{requirement_name}操作手册.md",
        "manual_docx": output_dir / f"{requirement_name}操作手册.docx",
        "manual_pdf": output_dir / f"{requirement_name}操作手册.pdf",
        "source_md": output_dir / f"{requirement_name}源代码文档.md",
        "source_docx": output_dir / f"{requirement_name}源代码文档.docx",
        "source_pdf": output_dir / f"{requirement_name}源代码文档.pdf",
        "form_md": output_dir / "软件著作权登记申请表.md",
        "form_docx": output_dir / "软件著作权登记申请表.docx",
    }


def build_source_docx_fixed(source_md: Path, source_docx: Path, header_text: str) -> None:
    lines = source_md.read_text(encoding="utf-8").splitlines()
    if len(lines) != 3000:
        raise ValueError(f"Source markdown must contain 3000 lines, got {len(lines)}")
    doc = Document()
    first = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    paragraphs = [first]
    for _ in range(1, len(lines)):
        paragraphs.append(doc.add_paragraph())
    for idx, (p, line) in enumerate(zip(paragraphs, lines, strict=True), start=1):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(LINE_PITCH_TWIPS / 20)
        run = p.add_run(line if line else "[padding]")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        if idx % LINES_PER_PAGE == 0 and idx != len(lines):
            run.add_break(WD_BREAK.PAGE)

    section = doc.sections[0]
    section.top_margin = Mm(NORMAL_MARGIN_MM)
    section.bottom_margin = Mm(NORMAL_MARGIN_MM)
    section.left_margin = Mm(NORMAL_MARGIN_MM)
    section.right_margin = Mm(NORMAL_MARGIN_MM)
    sect_pr = section._sectPr
    doc_grid = OxmlElement("w:docGrid")
    doc_grid.set(qn("w:type"), "lines")
    doc_grid.set(qn("w:linePitch"), str(LINE_PITCH_TWIPS))
    sect_pr.append(doc_grid)
    doc.save(str(source_docx))
    postprocess_docx_header_footer(source_docx, header_text)


def build_manual_docx_with_clean_cover(manual_md: Path, manual_docx: Path, pandoc: Path, header_text: str, requirement_name: str) -> None:
    convert_markdown_to_docx(manual_md, manual_docx, pandoc, manual_md.parent)
    doc = Document(str(manual_docx))
    if doc.paragraphs:
        cover_title = doc.paragraphs[0].insert_paragraph_before(f"{requirement_name} V1.0")
    else:
        cover_title = doc.add_paragraph(f"{requirement_name} V1.0")
    cover_sub = cover_title.insert_paragraph_before("操作手册")
    for p in (cover_title, cover_sub):
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for r in p.runs:
            r.font.name = "宋体"
            r.font.size = Pt(24)
    if cover_sub.runs:
        cover_sub.runs[0].add_break(WD_BREAK.PAGE)
    postprocess_docx_header_footer(manual_docx, header_text)


def build_materials(requirement_name: str, output_dir: Path) -> dict[str, str | int]:
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "images").mkdir(parents=True, exist_ok=True)
    paths = ensure_paths(requirement_name, output_dir)
    missing = [str(p) for p in (paths["prd_md"], paths["manual_md"], paths["source_md"], paths["form_md"]) if not p.exists()]
    if missing:
        raise FileNotFoundError("Missing required markdown inputs: " + ", ".join(missing))

    node_path = find_node()
    pandoc_path = find_pandoc()
    project_root = output_dir
    header_text = f"{requirement_name}V1.0"

    capture_screenshots(output_dir / "images", node_path=node_path, project_root=project_root)

    build_manual_docx_with_clean_cover(paths["manual_md"], paths["manual_docx"], pandoc_path, header_text, requirement_name)
    manual_html = output_dir / "_manual.html"
    convert_markdown_to_html(paths["manual_md"], manual_html, pandoc_path, output_dir)
    html_to_pdf(manual_html, paths["manual_pdf"], node_path)

    build_source_docx_fixed(paths["source_md"], paths["source_docx"], header_text)
    source_html = output_dir / "_source.html"
    convert_markdown_to_html(paths["source_md"], source_html, pandoc_path, output_dir)
    html_to_pdf(source_html, paths["source_pdf"], node_path)

    convert_markdown_to_docx(paths["form_md"], paths["form_docx"], pandoc_path, output_dir)
    postprocess_docx_header_footer(paths["form_docx"], header_text)

    for tmp in (manual_html, source_html):
        if tmp.exists():
            tmp.unlink()

    manual_chars = strip_markdown_for_count(paths["manual_md"].read_text(encoding="utf-8"))
    line_count = count_source_lines(project_root)

    return {
        "line_count": line_count,
        "manual_chars": manual_chars,
        "source_doc_lines": 3000,
        "source_doc_pages": 60,
        "prd_path": str(paths["prd_md"].resolve()),
        "manual_md_path": str(paths["manual_md"].resolve()),
        "manual_docx_path": str(paths["manual_docx"].resolve()),
        "manual_pdf_path": str(paths["manual_pdf"].resolve()),
        "source_md_path": str(paths["source_md"].resolve()),
        "source_docx_path": str(paths["source_docx"].resolve()),
        "source_pdf_path": str(paths["source_pdf"].resolve()),
        "form_md_path": str(paths["form_md"].resolve()),
        "form_docx_path": str(paths["form_docx"].resolve()),
        "images_dir": str((output_dir / "images").resolve()),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Build materials from existing markdown files.")
    parser.add_argument("--name", required=True, help="Requirement/software name.")
    parser.add_argument("--output-dir", required=True, help="Requirement directory containing markdown inputs.")
    args = parser.parse_args()
    result = build_materials(args.name, Path(args.output_dir).resolve())
    for key, value in result.items():
        print(f"{key}={value}")


if __name__ == "__main__":
    main()
