from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOFTWARE_NAME = "多维度绩效评价与智能分析系统"
ROOT = Path(__file__).resolve().parents[1]
MANUAL_MD = ROOT / f"{SOFTWARE_NAME}操作手册.md"
MANUAL_DOCX = ROOT / f"{SOFTWARE_NAME}操作手册.docx"
SOURCE_MD = ROOT / f"{SOFTWARE_NAME}源代码文档.md"
SOURCE_DOCX = ROOT / f"{SOFTWARE_NAME}源代码文档.docx"
FORM_MD = ROOT / "软件著作权登记申请表.md"
FORM_DOCX = ROOT / "软件著作权登记申请表.docx"

IMAGE_PATTERN = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)\s*$")
ORDERED_PATTERN = re.compile(r"^\d+\.\s+(?P<text>.*)$")
SOURCE_SUFFIXES = {".py", ".js", ".css", ".html", ".sql", ".json"}
EXCLUDE_DIRS = {".git", "node_modules", "__pycache__", "images"}


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_run_font(run, *, font_name: str = "宋体", size: int = 11, bold: bool = False) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def apply_doc_header_footer(document: Document, header_text: str) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    clear_paragraph(header)
    header_run = header.add_run(header_text)
    set_run_font(header_run, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    clear_paragraph(footer)
    add_page_number(footer)


def add_manual_cover(document: Document, title: str) -> None:
    cover = document.add_paragraph()
    cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    blank_run = cover.add_run("\n" * 4)
    set_run_font(blank_run, font_name="黑体", size=12)
    title_run = cover.add_run(title)
    set_run_font(title_run, font_name="黑体", size=22, bold=True)
    subtitle_run = cover.add_run("\n\n软件操作手册\n")
    set_run_font(subtitle_run, font_name="黑体", size=16, bold=False)
    version_run = cover.add_run("\n版本：V1.0\n\n")
    set_run_font(version_run, font_name="黑体", size=14, bold=False)
    date_run = cover.add_run("生成日期：2026-05-13")
    set_run_font(date_run, font_name="黑体", size=12, bold=False)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def convert_markdown_to_docx(md_path: Path, docx_path: Path, *, with_images: bool, add_cover: bool) -> None:
    document = Document()
    if add_cover:
        add_manual_cover(document, SOFTWARE_NAME)
    base_dir = md_path.parent
    ordered_index = 0

    for raw_line in md_path.read_text(encoding="utf-8").splitlines():
        stripped = raw_line.strip()
        if not stripped:
            document.add_paragraph("")
            ordered_index = 0
            continue

        if stripped.startswith("# "):
            paragraph = document.add_heading(level=1)
            set_run_font(paragraph.add_run(stripped[2:].strip()), font_name="黑体", size=16, bold=True)
            ordered_index = 0
            continue
        if stripped.startswith("## "):
            paragraph = document.add_heading(level=2)
            set_run_font(paragraph.add_run(stripped[3:].strip()), font_name="黑体", size=14, bold=True)
            ordered_index = 0
            continue
        if stripped.startswith("### "):
            paragraph = document.add_heading(level=3)
            set_run_font(paragraph.add_run(stripped[4:].strip()), font_name="黑体", size=12, bold=True)
            ordered_index = 0
            continue

        image_match = IMAGE_PATTERN.match(stripped)
        if with_images and image_match:
            alt = image_match.group("alt").strip()
            rel_path = image_match.group("path").strip()
            image_path = (base_dir / rel_path).resolve()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            if image_path.exists():
                run.add_picture(str(image_path), width=Inches(6.2))
            else:
                set_run_font(run)
                run.text = f"[图片缺失] {rel_path}"
            if alt:
                caption = document.add_paragraph()
                caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                set_run_font(caption.add_run(alt))
            ordered_index = 0
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            paragraph = document.add_paragraph(style="List Bullet")
            set_run_font(paragraph.add_run(stripped[2:].strip()))
            ordered_index = 0
            continue

        ordered_match = ORDERED_PATTERN.match(stripped)
        if ordered_match:
            ordered_index += 1
            paragraph = document.add_paragraph(style="List Number")
            set_run_font(paragraph.add_run(ordered_match.group("text").strip()))
            continue

        ordered_index = 0
        paragraph = document.add_paragraph()
        set_run_font(paragraph.add_run(raw_line.rstrip()))

    apply_doc_header_footer(document, f"{SOFTWARE_NAME}V1.0")
    document.save(docx_path)


def iter_source_files() -> list[Path]:
    files: list[Path] = []
    for path in ROOT.rglob("*"):
        if path.is_dir():
            continue
        if any(part in EXCLUDE_DIRS for part in path.relative_to(ROOT).parts):
            continue
        if path.suffix.lower() not in SOURCE_SUFFIXES:
            continue
        files.append(path)
    files.sort(key=lambda item: item.relative_to(ROOT).as_posix())
    return files


def build_source_markdown() -> list[str]:
    lines: list[str] = []
    for file_path in iter_source_files():
        relative = file_path.relative_to(ROOT).as_posix()
        lines.append(f"文件：{relative}")
        for index, raw_line in enumerate(file_path.read_text(encoding="utf-8", errors="ignore").splitlines(), start=1):
            lines.append(f"{index:04d} {raw_line}")
        lines.append("")
    if len(lines) < 3000:
        lines.extend(["0000 "] * (3000 - len(lines)))
    return lines[:3000]


def write_source_documents() -> None:
    source_lines = build_source_markdown()
    SOURCE_MD.write_text("\n".join(source_lines), encoding="utf-8")

    document = Document()
    apply_doc_header_footer(document, f"{SOFTWARE_NAME}V1.0")
    for index, line in enumerate(source_lines, start=1):
      paragraph = document.add_paragraph()
      run = paragraph.add_run(line if line else " ")
      set_run_font(run, font_name="Consolas", size=8)
      paragraph.paragraph_format.space_before = 0
      paragraph.paragraph_format.space_after = 0
      paragraph.paragraph_format.line_spacing = 1.0
      if index % 50 == 0 and index != len(source_lines):
          run.add_break(WD_BREAK.PAGE)
    document.save(SOURCE_DOCX)


def main() -> None:
    write_source_documents()
    convert_markdown_to_docx(MANUAL_MD, MANUAL_DOCX, with_images=True, add_cover=True)
    convert_markdown_to_docx(FORM_MD, FORM_DOCX, with_images=False, add_cover=False)
    print(MANUAL_DOCX)
    print(SOURCE_DOCX)
    print(FORM_DOCX)


if __name__ == "__main__":
    main()
