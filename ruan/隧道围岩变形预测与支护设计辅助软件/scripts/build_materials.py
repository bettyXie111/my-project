from __future__ import annotations

import argparse
import html
import os
import re
import shutil
import subprocess
import sys
import time
from pathlib import Path
from urllib.error import URLError
from urllib.request import urlopen

from docx import Document
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
EDGE_EXECUTABLE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")
DEFAULT_SERVER_PORT = 8010
CODE_EXTENSIONS = {".py", ".js", ".mjs", ".css", ".html", ".json"}
TOP_BOTTOM_MARGIN_MM = 18
LEFT_RIGHT_MARGIN_MM = 14
NORMAL_MARGIN_MM = 25.4
EDITABLE_LINES_PER_PAGE = 50
EDITABLE_LINE_PITCH_PT = ((297 - NORMAL_MARGIN_MM * 2) / EDITABLE_LINES_PER_PAGE) * 72 / 25.4
EDITABLE_LINE_PITCH_TWIPS = round(EDITABLE_LINE_PITCH_PT * 20)
PAGE_HEIGHT_TWIPS = 16840
EDITABLE_TOP_BOTTOM_MARGIN_TWIPS = (PAGE_HEIGHT_TWIPS - EDITABLE_LINE_PITCH_TWIPS * EDITABLE_LINES_PER_PAGE) // 2
EDITABLE_GRID_RESIDUAL_TWIPS = PAGE_HEIGHT_TWIPS - (EDITABLE_TOP_BOTTOM_MARGIN_TWIPS * 2 + EDITABLE_LINE_PITCH_TWIPS * EDITABLE_LINES_PER_PAGE)


def find_workspace_root(start: Path) -> Path:
    for parent in [start, *start.parents]:
        if (parent / ".tools").exists():
            return parent
    raise FileNotFoundError("Cannot locate workspace root containing .tools")


def find_skill_root(workspace_root: Path) -> Path:
    candidates = [
        workspace_root / ".external_skills" / "chinese-copyright-application",
        Path.home() / ".codex" / "skills" / "custom" / "chinese-copyright-application",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    raise FileNotFoundError("Cannot locate chinese-copyright-application skill directory")


WORKSPACE_ROOT = find_workspace_root(PROJECT_ROOT)
SKILL_ROOT = find_skill_root(WORKSPACE_ROOT)
MANUAL_REFERENCE_DOCX = SKILL_ROOT / "references" / "operation-manual-template.docx"


def run_command(command: list[str], *, cwd: Path = PROJECT_ROOT, env: dict[str, str] | None = None) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        command,
        cwd=cwd,
        env=env,
        check=True,
        text=True,
        capture_output=True,
        encoding="utf-8",
        errors="ignore",
    )


def find_first_file(base: Path, pattern: str) -> Path:
    matches = sorted(base.rglob(pattern))
    if not matches:
        raise FileNotFoundError(f"Cannot find {pattern} under {base}")
    return matches[-1]


def find_node() -> Path:
    return find_first_file(WORKSPACE_ROOT / ".tools" / "node", "node.exe")


def find_pandoc() -> Path:
    return find_first_file(WORKSPACE_ROOT / ".tools" / "pandoc", "pandoc.exe")


def iter_source_files(base_dir: Path) -> list[Path]:
    files: list[Path] = []
    for root_name in ("apps", "packages", "scripts"):
        base = base_dir / root_name
        if not base.exists():
            continue
        for path in sorted(base.rglob("*")):
            if path.is_file() and path.suffix.lower() in CODE_EXTENSIONS:
                if "__pycache__" in path.parts or "node_modules" in path.parts:
                    continue
                files.append(path)
    return files


def count_source_lines(base_dir: Path) -> int:
    total = 0
    for path in iter_source_files(base_dir):
        total += sum(1 for line in path.read_text(encoding="utf-8", errors="ignore").splitlines() if line.strip())
    return total


def strip_markdown_for_count(markdown_text: str) -> int:
    visible = []
    for line in markdown_text.splitlines():
        if line.strip().startswith("!["):
            continue
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", line)
        visible.append(line.replace("`", ""))
    return len(re.sub(r"\s+", "", "".join(visible)))


def ensure_paths(requirement_name: str, output_dir: Path) -> dict[str, Path]:
    return {
        "prd_md": output_dir / f"{requirement_name}需求规格说明书.md",
        "manual_md": output_dir / f"{requirement_name}操作手册.md",
        "manual_docx": output_dir / f"{requirement_name}操作手册.docx",
        "manual_pdf": output_dir / f"{requirement_name}操作手册.pdf",
        "source_md": output_dir / f"{requirement_name}源代码文档.md",
        "source_docx": output_dir / f"{requirement_name}源代码文档.docx",
        "source_edit_docx": output_dir / f"{requirement_name}源代码文档-编辑版.docx",
        "source_pdf": output_dir / f"{requirement_name}源代码文档.pdf",
        "form_md": output_dir / "软件著作权登记申请表.md",
        "form_docx": output_dir / "软件著作权登记申请表.docx",
    }


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def set_run_fonts(run, western_font: str, east_asia_font: str, font_size: int, *, bold: bool = False) -> None:
    run.font.name = western_font
    run.font.size = Pt(font_size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), western_font)
    rfonts.set(qn("w:hAnsi"), western_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def clear_paragraph(paragraph) -> None:
    node = paragraph._element
    for child in list(node):
        if child.tag.endswith("}pPr"):
            continue
        node.remove(child)


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def normalize_all_section_headers(doc: Document, header_text: str) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        clear_paragraph(header)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header.add_run(header_text)
        set_run_fonts(header_run, "Times New Roman", "瀹嬩綋", 10)

        section.top_margin = Mm(TOP_BOTTOM_MARGIN_MM)
        section.bottom_margin = Mm(TOP_BOTTOM_MARGIN_MM)
        section.left_margin = Mm(LEFT_RIGHT_MARGIN_MM)
        section.right_margin = Mm(LEFT_RIGHT_MARGIN_MM)

        footer = section.footer.paragraphs[0]
        clear_paragraph(footer)
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        add_page_number(footer)


def apply_word_line_numbering(section, *, start: int = 1, count_by: int = 1, distance_twips: int = 360) -> None:
    sect_pr = section._sectPr
    for child in list(sect_pr):
        if child.tag == qn("w:lnNumType"):
            sect_pr.remove(child)
    ln_num = OxmlElement("w:lnNumType")
    ln_num.set(qn("w:countBy"), str(count_by))
    ln_num.set(qn("w:start"), str(start))
    ln_num.set(qn("w:distance"), str(distance_twips))
    ln_num.set(qn("w:restart"), "continuous")
    sect_pr.append(ln_num)


def apply_fixed_line_grid(section, *, line_pitch_twips: int = EDITABLE_LINE_PITCH_TWIPS) -> None:
    sect_pr = section._sectPr
    for child in list(sect_pr):
        if child.tag == qn("w:docGrid"):
            sect_pr.remove(child)
    doc_grid = OxmlElement("w:docGrid")
    doc_grid.set(qn("w:type"), "lines")
    doc_grid.set(qn("w:linePitch"), str(line_pitch_twips))
    sect_pr.append(doc_grid)


def apply_top_bottom_margins_twips(section, *, margin_twips: int) -> None:
    sect_pr = section._sectPr
    pg_mar = sect_pr.find(qn("w:pgMar"))
    if pg_mar is None:
        pg_mar = OxmlElement("w:pgMar")
        sect_pr.append(pg_mar)
    pg_mar.set(qn("w:top"), str(margin_twips))
    pg_mar.set(qn("w:bottom"), str(margin_twips))


def validate_editable_page_grid() -> None:
    if EDITABLE_GRID_RESIDUAL_TWIPS != 0:
        raise ValueError(
            "Editable DOCX page grid is not closed for 50 lines/page: "
            f"residual_twips={EDITABLE_GRID_RESIDUAL_TWIPS}, "
            f"line_pitch_twips={EDITABLE_LINE_PITCH_TWIPS}, "
            f"top_bottom_margin_twips={EDITABLE_TOP_BOTTOM_MARGIN_TWIPS}"
        )


def set_editable_code_paragraph_layout(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(EDITABLE_LINE_PITCH_PT)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    ppr = paragraph._element.get_or_add_pPr()
    for child in list(ppr):
        if child.tag == qn("w:snapToGrid"):
            ppr.remove(child)
    ppr.append(OxmlElement("w:snapToGrid"))


def convert_markdown_to_docx(
    md_path: Path,
    docx_path: Path,
    pandoc_path: Path,
    resource_path: Path | None = None,
    reference_doc: Path | None = None,
) -> None:
    cmd = [str(pandoc_path), str(md_path), "-f", "gfm", "-t", "docx", "-o", str(docx_path)]
    if resource_path is not None:
        cmd.extend(["--resource-path", str(resource_path)])
    if reference_doc is not None:
        cmd.extend(["--reference-doc", str(reference_doc)])
    run_command(cmd)


def convert_markdown_to_html(md_path: Path, html_path: Path, pandoc_path: Path, resource_path: Path | None = None) -> None:
    cmd = [str(pandoc_path), str(md_path), "-f", "gfm", "-t", "html5", "--standalone", "-o", str(html_path)]
    if resource_path is not None:
        cmd.extend(["--resource-path", str(resource_path)])
    run_command(cmd)


def html_to_pdf(html_path: Path, pdf_path: Path, node_path: Path, header_text: str) -> None:
    env = os.environ.copy()
    env["EDGE_EXECUTABLE_PATH"] = str(EDGE_EXECUTABLE)
    env["PDF_HEADER_TEXT"] = header_text
    run_command([str(node_path), str(PROJECT_ROOT / "scripts" / "html_to_pdf_with_header.mjs"), str(html_path), str(pdf_path)], env=env)


def wait_for_service(base_url: str, timeout_seconds: float = 20.0) -> None:
    deadline = time.time() + timeout_seconds
    while time.time() < deadline:
        try:
            with urlopen(base_url, timeout=2):
                return
        except URLError:
            time.sleep(0.5)
    raise TimeoutError(f"Service not available at {base_url}")


def capture_screenshots(images_dir: Path, node_path: Path) -> None:
    base_url = f"http://127.0.0.1:{DEFAULT_SERVER_PORT}"
    env = os.environ.copy()
    env["APP_PORT"] = str(DEFAULT_SERVER_PORT)
    server = subprocess.Popen([sys.executable, "apps/api/main.py"], cwd=PROJECT_ROOT, env=env)
    try:
        wait_for_service(base_url)
        shot_env = os.environ.copy()
        shot_env["SCREENSHOT_OUT_DIR"] = str(images_dir)
        shot_env["SCREENSHOT_BASE_URL"] = base_url
        shot_env["EDGE_EXECUTABLE_PATH"] = str(EDGE_EXECUTABLE)
        run_command([str(node_path), str(PROJECT_ROOT / "scripts" / "capture_screenshots_local.mjs")], env=shot_env)
    finally:
        server.terminate()
        server.wait(timeout=5)


def postprocess_operation_manual_docx(docx_path: Path, header_text: str, requirement_name: str) -> None:
    doc = Document(str(docx_path))
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        remove_paragraph(doc.paragraphs[0])
    if not doc.paragraphs:
        doc.add_paragraph("姝ｆ枃")
    first = doc.paragraphs[0]
    subtitle = first.insert_paragraph_before("鎿嶄綔鎵嬪唽")
    title = subtitle.insert_paragraph_before(f"{requirement_name} V1.0")

    for paragraph in (title, subtitle):
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_fonts(run, "Times New Roman", "瀹嬩綋", 24, bold=False)
    title.paragraph_format.space_before = Pt(220)
    subtitle.runs[0].add_break(WD_BREAK.PAGE)

    normalize_all_section_headers(doc, header_text)
    doc.save(str(docx_path))


def build_operation_manual_docx(md_path: Path, docx_path: Path, pandoc_path: Path, header_text: str, requirement_name: str) -> None:
    convert_markdown_to_docx(
        md_path,
        docx_path,
        pandoc_path,
        resource_path=md_path.parent,
        reference_doc=MANUAL_REFERENCE_DOCX,
    )
    postprocess_operation_manual_docx(docx_path, header_text, requirement_name)


def build_operation_manual_html(md_path: Path, html_path: Path, pandoc_path: Path, requirement_name: str) -> None:
    convert_markdown_to_html(md_path, html_path, pandoc_path, md_path.parent)
    original = html_path.read_text(encoding="utf-8")
    cover = f"""
<section class="cover-page">
  <div class="cover-inner">
    <div class="cover-title">{html.escape(requirement_name)} V1.0</div>
    <div class="cover-subtitle">鎿嶄綔鎵嬪唽</div>
  </div>
</section>
"""
    style = """
<style>
body { font-family: "SimSun", "瀹嬩綋", serif; color: #000; }
.cover-page {
  height: 260mm;
  display: flex;
  align-items: center;
  justify-content: center;
  page-break-after: always;
}
.cover-inner { text-align: center; }
.cover-title, .cover-subtitle {
  font-family: "SimSun", "瀹嬩綋", serif;
  font-size: 24pt;
  line-height: 1.35;
}
img { max-width: 100%; height: auto; }
p { line-height: 1.7; }
</style>
"""
    updated = original.replace("</head>", f"{style}</head>").replace("<body>", f"<body>{cover}", 1)
    html_path.write_text(updated, encoding="utf-8")


def build_source_docx(source_md_path: Path, source_docx_path: Path, header_text: str) -> None:
    validate_editable_page_grid()
    lines = source_md_path.read_text(encoding="utf-8").splitlines()
    if len(lines) != 3000:
        raise ValueError(f"Source markdown must contain 3000 lines, got {len(lines)}")

    doc = Document()
    if doc.paragraphs:
        first_paragraph = doc.paragraphs[0]
    else:
        first_paragraph = doc.add_paragraph()
    paragraphs = [first_paragraph]
    for _ in range(1, len(lines)):
        paragraphs.append(doc.add_paragraph())

    for paragraph, line in zip(paragraphs, lines, strict=True):
        clear_paragraph(paragraph)
        set_editable_code_paragraph_layout(paragraph)
        run = paragraph.add_run(line or " ")
        set_run_fonts(run, "Consolas", "新宋体", 10)

    normalize_all_section_headers(doc, header_text)
    for section in doc.sections:
        section.top_margin = Mm(NORMAL_MARGIN_MM)
        section.bottom_margin = Mm(NORMAL_MARGIN_MM)
        section.left_margin = Mm(NORMAL_MARGIN_MM)
        section.right_margin = Mm(NORMAL_MARGIN_MM)
        apply_fixed_line_grid(section)
        apply_top_bottom_margins_twips(section, margin_twips=EDITABLE_TOP_BOTTOM_MARGIN_TWIPS)
    doc.save(str(source_docx_path))


def strip_submit_line_number(line: str) -> str:
    parts = line.split(" " * 10, 1)
    if len(parts) == 2 and parts[0].isdigit():
        return parts[1]
    return line


def build_source_edit_docx(source_md_path: Path, source_edit_docx_path: Path, header_text: str) -> None:
    validate_editable_page_grid()
    numbered_lines = source_md_path.read_text(encoding="utf-8").splitlines()
    if len(numbered_lines) != 3000:
        raise ValueError(f"Source markdown must contain 3000 lines, got {len(numbered_lines)}")

    doc = Document()
    if doc.paragraphs:
        first_paragraph = doc.paragraphs[0]
    else:
        first_paragraph = doc.add_paragraph()
    paragraphs = [first_paragraph]
    for _ in range(1, len(numbered_lines)):
        paragraphs.append(doc.add_paragraph())

    for paragraph, line in zip(paragraphs, numbered_lines, strict=True):
        clear_paragraph(paragraph)
        set_editable_code_paragraph_layout(paragraph)
        run = paragraph.add_run(strip_submit_line_number(line) or " ")
        set_run_fonts(run, "Consolas", "新宋体", 10)

    normalize_all_section_headers(doc, header_text)
    for section in doc.sections:
        # Word may consume a hidden top grid line under normal margins, which shifts
        # the first visible code line to "2" when start=1. Use start=0 as compensation
        # so the first visible line number is rendered as "1".
        section.top_margin = Mm(NORMAL_MARGIN_MM)
        section.bottom_margin = Mm(NORMAL_MARGIN_MM)
        section.left_margin = Mm(NORMAL_MARGIN_MM)
        section.right_margin = Mm(NORMAL_MARGIN_MM)
        apply_fixed_line_grid(section)
        apply_top_bottom_margins_twips(section, margin_twips=EDITABLE_TOP_BOTTOM_MARGIN_TWIPS)
    doc.save(str(source_edit_docx_path))


def build_source_html(source_md_path: Path, html_path: Path) -> None:
    lines = source_md_path.read_text(encoding="utf-8").splitlines()
    if len(lines) != 3000:
        raise ValueError(f"Source markdown must contain 3000 lines, got {len(lines)}")

    pages = []
    for page_index in range(0, len(lines), 50):
        chunk = lines[page_index : page_index + 50]
        rows = "\n".join(f"<div class=\"code-line\">{html.escape(line)}</div>" for line in chunk)
        pages.append(f"<section class=\"page\">{rows}</section>")

    document = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <title>婧愪唬鐮佹枃妗?/title>
  <style>
    @page {{ size: A4; margin: 18mm 14mm; }}
    body {{ margin: 0; font-family: "Consolas", "Courier New", "NSimSun", monospace; color: #000; }}
    .page {{ page-break-after: always; min-height: 255mm; }}
    .page:last-child {{ page-break-after: auto; }}
    .code-line {{ white-space: pre; font-size: 10pt; line-height: 1.26; }}
  </style>
</head>
<body>
{''.join(pages)}
</body>
</html>"""
    html_path.write_text(document, encoding="utf-8")


def postprocess_application_form_docx(docx_path: Path, header_text: str) -> None:
    doc = Document(str(docx_path))
    normalize_all_section_headers(doc, header_text)
    doc.save(str(docx_path))


def build_materials(requirement_name: str, output_dir: Path) -> dict[str, str | int]:
    output_dir.mkdir(parents=True, exist_ok=True)
    images_dir = output_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    run_command([sys.executable, "scripts/build_source_doc.py"], cwd=PROJECT_ROOT)

    paths = ensure_paths(requirement_name, output_dir)
    missing = [str(p) for p in (paths["prd_md"], paths["manual_md"], paths["source_md"], paths["form_md"]) if not p.exists()]
    if missing:
        raise FileNotFoundError("Missing required markdown inputs: " + ", ".join(missing))

    header_text = f"{requirement_name}V1.0"
    node_path = find_node()
    pandoc_path = find_pandoc()
    capture_screenshots(images_dir, node_path=node_path)

    build_operation_manual_docx(paths["manual_md"], paths["manual_docx"], pandoc_path, header_text, requirement_name)
    manual_html = output_dir / "_manual.html"
    build_operation_manual_html(paths["manual_md"], manual_html, pandoc_path, requirement_name)
    html_to_pdf(manual_html, paths["manual_pdf"], node_path, header_text)

    build_source_docx(paths["source_md"], paths["source_docx"], header_text)
    build_source_edit_docx(paths["source_md"], paths["source_edit_docx"], header_text)
    source_html = output_dir / "_source.html"
    build_source_html(paths["source_md"], source_html)
    html_to_pdf(source_html, paths["source_pdf"], node_path, header_text)

    convert_markdown_to_docx(paths["form_md"], paths["form_docx"], pandoc_path, output_dir)
    postprocess_application_form_docx(paths["form_docx"], header_text)

    for tmp in (manual_html, source_html):
        if tmp.exists():
            tmp.unlink()

    manual_chars = strip_markdown_for_count(paths["manual_md"].read_text(encoding="utf-8"))
    line_count = count_source_lines(output_dir)
    source_line_count = len(paths["source_md"].read_text(encoding="utf-8").splitlines())

    return {
        "line_count": line_count,
        "manual_chars": manual_chars,
        "source_doc_lines": source_line_count,
        "source_doc_pages": 60,
        "prd_path": str(paths["prd_md"].resolve()),
        "manual_md_path": str(paths["manual_md"].resolve()),
        "manual_docx_path": str(paths["manual_docx"].resolve()),
        "manual_pdf_path": str(paths["manual_pdf"].resolve()),
        "source_md_path": str(paths["source_md"].resolve()),
        "source_docx_path": str(paths["source_docx"].resolve()),
        "source_edit_docx_path": str(paths["source_edit_docx"].resolve()),
        "source_pdf_path": str(paths["source_pdf"].resolve()),
        "form_md_path": str(paths["form_md"].resolve()),
        "form_docx_path": str(paths["form_docx"].resolve()),
        "images_dir": str(images_dir.resolve()),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Build copyright materials from project markdown files.")
    parser.add_argument("--name", required=True, help="Requirement/software name.")
    parser.add_argument("--output-dir", required=True, help="Requirement directory containing markdown inputs.")
    args = parser.parse_args()
    result = build_materials(args.name, Path(args.output_dir).resolve())
    for key, value in result.items():
        print(f"{key}={value}")


if __name__ == "__main__":
    main()



