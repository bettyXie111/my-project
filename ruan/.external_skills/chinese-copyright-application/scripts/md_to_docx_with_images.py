from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

IMAGE_PATTERN = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)\s*$")
ORDERED_PATTERN = re.compile(r"^\d+\.\s+(?P<text>.*)$")


def set_run_black(run) -> None:
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")


def add_text_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_black(r)


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    base_dir = md_path.parent
    ordered_index = 0

    for raw_line in md_path.read_text(encoding="utf-8").splitlines():
        stripped = raw_line.strip()

        if not stripped:
            doc.add_paragraph("")
            ordered_index = 0
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(level=1)
            set_run_black(p.add_run(stripped[2:].strip()))
            ordered_index = 0
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(level=2)
            set_run_black(p.add_run(stripped[3:].strip()))
            ordered_index = 0
            continue
        if stripped.startswith("### "):
            p = doc.add_heading(level=3)
            set_run_black(p.add_run(stripped[4:].strip()))
            ordered_index = 0
            continue

        image_match = IMAGE_PATTERN.match(stripped)
        if image_match:
            alt = image_match.group("alt").strip()
            rel_path = image_match.group("path").strip()
            image_path = (base_dir / rel_path).resolve()
            if image_path.exists():
                p = doc.add_paragraph()
                p.alignment = 1
                r = p.add_run()
                r.add_picture(str(image_path), width=Inches(6.2))
                if alt:
                    add_text_paragraph(doc, alt)
            else:
                add_text_paragraph(doc, f"[图片缺失] {rel_path}")
            ordered_index = 0
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            add_text_paragraph(doc, f"• {stripped[2:].strip()}")
            ordered_index = 0
            continue

        ordered = ORDERED_PATTERN.match(stripped)
        if ordered:
            ordered_index += 1
            add_text_paragraph(doc, f"{ordered_index}. {ordered.group('text').strip()}")
            continue

        ordered_index = 0
        add_text_paragraph(doc, raw_line.rstrip())

    doc.save(str(docx_path))


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 3:
        raise SystemExit("Usage: python md_to_docx_with_images.py <input.md> <output.docx>")
    convert(Path(sys.argv[1]).resolve(), Path(sys.argv[2]).resolve())
    print(f"Generated: {Path(sys.argv[2]).resolve()}")
