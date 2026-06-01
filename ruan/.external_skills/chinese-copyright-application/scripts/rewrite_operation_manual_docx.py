from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_block_items(parent: _Document):
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        else:
            yield Table(child, parent)


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def set_paragraph(paragraph: Paragraph, text: str, style: str | None = None) -> None:
    clear_paragraph(paragraph)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)


def main() -> None:
    docx_files = sorted(Path(".").glob("*.docx"))
    if not docx_files:
        raise FileNotFoundError("No .docx file found in current directory.")

    path = docx_files[0]
    doc = Document(str(path))

    template_lines: list[tuple[str, str]] = [
        ("Heading 1", "智路城市道路养护管理平台软件 V1.0"),
        ("Normal", "操作手册模板"),
        ("Normal", "本模板仅保留版式结构、目录层级和示例占位，不保留大段业务说明正文。"),
        ("Heading 2", "一、系统入口与登录"),
        ("Heading 3", "1.1 登录界面总览"),
        ("Normal", "用于放置登录界面完整截图，并展示系统名称、标题、导航及输入区域。"),
        ("Heading 3", "1.2 登录操作步骤"),
        ("List Number", "打开系统入口，进入登录页面。"),
        ("List Number", "输入账号、密码等必要信息。"),
        ("List Number", "点击“登录”，跳转至首页。"),
        ("List Number", "记录登录结果及关键状态变化。"),
        ("Heading 2", "二、首页操作"),
        ("Heading 3", "2.1 首页布局与导航"),
        ("Normal", "用于放置首页完整截图，需展示顶部栏、左侧功能栏、标题及核心内容区。"),
        ("Heading 3", "2.2 首页关键流程"),
        ("List Number", "查看首页关键指标或待办信息。"),
        ("List Number", "点击一级菜单进入目标模块。"),
        ("List Number", "记录页面跳转及状态变化。"),
        ("Heading 2", "三、一级功能模块与二级菜单操作"),
        ("Normal", "要求：一级功能介绍完整；需展示二级菜单展开、下拉框展开、复选框展开等操作状态。"),
        ("Heading 3", "3.1 一级功能A"),
        ("Normal", "功能说明占位。"),
        ("Heading 3", "3.2 一级功能A-二级菜单展开"),
        ("List Number", "点击一级菜单，展开全部二级菜单项。"),
        ("List Number", "逐项进入对应页面并说明跳转关系。"),
        ("Heading 3", "3.3 一级功能A-下拉框与复选框操作"),
        ("List Number", "展开下拉框，展示全部可选项。"),
        ("List Number", "勾选复选框并执行查询、提交或批量处理。"),
        ("List Number", "记录数据变化和结果反馈。"),
        ("Heading 3", "3.4 一级功能B"),
        ("Normal", "按相同结构补充功能说明、二级菜单展开和操作步骤。"),
        ("Heading 3", "3.5 其余一级功能"),
        ("Normal", "逐个一级功能完整覆盖；当手册少于30页时，全部功能点均需展开说明。"),
        ("Heading 2", "四、端到端操作流程演示"),
        ("Heading 3", "4.1 业务流程主线"),
        ("List Number", "登录系统。"),
        ("List Number", "进入首页。"),
        ("List Number", "从一级菜单进入二级页面。"),
        ("List Number", "完成查询、编辑、提交等关键动作。"),
        ("List Number", "返回首页查看结果或统计变化。"),
        ("Heading 3", "4.2 页面跳转与数据变化追踪"),
        ("Normal", "起点页面：占位"),
        ("Normal", "操作动作：占位"),
        ("Normal", "目标页面：占位"),
        ("Normal", "数据变化：占位"),
        ("Heading 2", "五、附录"),
        ("Heading 3", "5.1 常见问题"),
        ("Normal", "问题与解决方案占位。"),
        ("Heading 3", "5.2 版本信息"),
        ("Normal", "软件全称：占位"),
        ("Normal", "版本号：V1.0"),
        ("Normal", "文档版本：占位"),
        ("Normal", "更新日期：占位"),
    ]

    paragraphs = doc.paragraphs
    if not paragraphs:
        raise ValueError("The document contains no paragraphs.")

    for index, (style, text) in enumerate(template_lines):
        if index < len(paragraphs):
            set_paragraph(paragraphs[index], text, style)
        else:
            p = doc.add_paragraph(style=style)
            p.add_run(text)

    for paragraph in paragraphs[len(template_lines):]:
        set_paragraph(paragraph, "", None)

    body = doc.element.body
    for block in list(iter_block_items(doc)):
        if isinstance(block, Table):
            body.remove(block._element)

    doc.save(str(path))
    print(f"Updated: {path}")


if __name__ == "__main__":
    main()
